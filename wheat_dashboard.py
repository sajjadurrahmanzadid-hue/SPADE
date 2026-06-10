"""
SPADE — Statistical Platform for Agronomic Data Evaluation
Dhaka University Nanotechnology Centre
Run with: streamlit run wheat_dashboard.py
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from scipy import stats
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from statsmodels.formula.api import ols
from statsmodels.stats.anova import anova_lm
from itertools import combinations
import os, io, warnings
warnings.filterwarnings("ignore")

# ── PAGE CONFIG ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="SPADE — Statistical Platform for Agronomic Data Evaluation",
    page_icon="🌿",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
  [data-testid="stMetricValue"] { font-size: 1.1rem; font-weight: 700 }
  .block-container { padding-top: 1.2rem }
  thead tr th { background: #f0f2f6 !important; font-size: 12px }
  .stTabs [data-baseweb="tab"] { font-size: 13px; font-weight: 600 }
  .sig-star { color: #c0392b; font-weight: 800 }
</style>""", unsafe_allow_html=True)

# ── CONSTANTS ─────────────────────────────────────────────────────────────────
TREATMENTS = [
    # n = total N actually applied (RDN + proportional topdress urea)
    # Topdress (87 kg urea/ha, proportional to RDN level) added at CRI stage T2–T8
    # 100% RDN: 130 + 45   = 175.00 kg N/ha
    #  75% RDN: 97.5 + 33.75 = 131.25 kg N/ha
    #  50% RDN: 65   + 22.5  = 87.50  kg N/ha
    {"id":"T1","n":0,      "type":"control","desc":"Absolute control (N₀P₀K₀)"},
    {"id":"T2","n":175,    "type":"water",  "desc":"100% RDN (175 kg N/ha total) + Water spray"},
    {"id":"T3","n":175,    "type":"nano",   "desc":"100% RDN (175 kg N/ha total) + Nano urea"},
    {"id":"T4","n":131.25, "type":"nano",   "desc":"75% RDN (131.25 kg N/ha total) + Nano urea"},
    {"id":"T5","n":131.25, "type":"water",  "desc":"75% RDN (131.25 kg N/ha total) + Water spray"},
    {"id":"T6","n":87.5,   "type":"water",  "desc":"50% RDN (87.5 kg N/ha total) + Water spray"},
    {"id":"T7","n":87.5,   "type":"nano",   "desc":"50% RDN (87.5 kg N/ha total) + Nano urea"},
    {"id":"T8","n":87.5,   "type":"gran",   "desc":"50% RDN (87.5 kg N/ha total) + Granulated urea spray"},
]
T_IDS   = [t["id"] for t in TREATMENTS]
T_INFO  = {t["id"]: t for t in TREATMENTS}
# Okabe-Ito colorblind-safe palette (Wong 2011, Nature Methods)
T_COLORS = {
    "control": "#808080",   # neutral grey
    "water":   "#0072B2",   # deep blue
    "nano":    "#009E73",   # teal green
    "gran":    "#E69F00",   # amber
}
T_COLORS_LIGHT = {
    "control": "#C8C8C8",
    "water":   "#80C4E4",
    "nano":    "#80D0B9",
    "gran":    "#F5CF80",
}

# ── 2-WAY ANOVA FACTOR MAPPINGS ───────────────────────────────────────────────
FACTOR_N = {   # Factor A: N rate level
    "T1":"N0 (0%)","T2":"N3 (100%)","T3":"N3 (100%)",
    "T4":"N2 (75%)","T5":"N2 (75%)","T6":"N1 (50%)",
    "T7":"N1 (50%)","T8":"N1 (50%)",
}
FACTOR_F = {   # Factor B: Foliar application type
    "T1":"Control","T2":"Water","T3":"Nano",
    "T4":"Nano","T5":"Water","T6":"Water",
    "T7":"Nano","T8":"Gran",
}
BALANCED_TIDS = ["T2","T3","T4","T5","T6","T7"]  # 3×2 balanced subset

PARAM_GROUPS = {
    "Growth & Biometrics": [
        ("plantH",      "Plant Height (cm)",              1),
        ("tillerN",     "Tiller Count",                   0),
        ("spad_t",      "SPAD — Tillering",               1),
        ("spad_h",      "SPAD — Heading",                 1),
        ("lai",         "Leaf Area Index",                2),
        ("stemD",       "Stem Diameter (mm)",             2),
        ("spikeL",      "Spike Length (cm)",              2),
        ("rootL",       "Root Length (cm)",               1),
        ("shootWt",     "Shoot Fresh Weight (g/pot)",     2),
        ("rootWt",      "Root Fresh Weight (g/pot)",      2),
        ("shootDW",     "Shoot Dry Weight (g/pot)",       2),
        ("rootDW",      "Root Dry Weight (g/pot)",        2),
    ],
    "Harvest & Quality": [
        ("grainsSpk",   "Grains per Spike",              1),
        ("spikletsSpk", "Spikelets per Spike",           1),
        ("tgw",         "1000-Grain Weight (g)",         2),
        ("grainY",      "Grain Yield (g/pot)",           2),
        ("strawY",      "Straw Yield (g/pot)",           2),
        ("grainN",      "Grain N (%)",                   2),
        ("strawN",      "Straw N (%)",                   2),
        ("grainB",      "Grain B (mg/kg)",               2),
        ("strawB",      "Straw B (mg/kg)",               2),
        ("grainP",      "Grain P (mg/kg)",               1),
        ("strawP",      "Straw P (mg/kg)",               1),
        ("grainK",      "Grain K (mg/kg)",               1),
        ("strawK",      "Straw K (mg/kg)",               1),
        ("grainS",      "Grain S (mg/kg)",               1),
        ("strawS",      "Straw S (mg/kg)",               1),
        ("grainCa",     "Grain Ca (mg/kg)",              1),
        ("strawCa",     "Straw Ca (mg/kg)",              1),
    ],
    "Soil Analysis": [
        ("soilPH",      "Soil pH",                       2),
        ("soilEC",      "Soil EC (dS/m)",                3),
        ("soilOC",      "Organic C (%)",                 2),
        ("totalN",      "Total Kjeldahl N (mg/kg)",      1),
        ("residN",      "Residual Mineral N (mg/kg)",    1),
        ("availP",      "Available P (mg/kg)",           2),
        ("excK",        "Exch. K (cmol/kg)",             3),
        ("availS",      "Available S (mg/kg)",           2),
        ("excCa",       "Exch. Ca (cmol/kg)",            2),
        ("hwB",         "Hot-water B (mg/kg)",           3),
    ],
}
ALL_PARAMS = [(k,l,d) for g in PARAM_GROUPS.values() for k,l,d in g]
P_LABEL = {k:l for k,l,d in ALL_PARAMS}
P_DEC   = {k:d for k,l,d in ALL_PARAMS}
P_KEYS  = [k for k,l,d in ALL_PARAMS]
# Always save next to the .py file, regardless of where terminal was launched
DATA_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "wheat_data.xlsx")

# ── DATA INIT ─────────────────────────────────────────────────────────────────
def blank_df():
    rows = []
    for t in TREATMENTS:
        for r in ["R1","R2","R3"]:
            row = {"Treatment": t["id"], "Replicate": r}
            for k in P_KEYS:
                row[k] = np.nan
            rows.append(row)
    return pd.DataFrame(rows)

def load_data():
    if os.path.exists(DATA_FILE):
        try:
            # Try named-sheet format first (v2+), fall back to flat file
            try:
                df = pd.read_excel(DATA_FILE, sheet_name="Raw Data")
            except Exception:
                df = pd.read_excel(DATA_FILE)
            base = blank_df()
            for col in base.columns:
                if col not in df.columns:
                    df[col] = np.nan
            return df[base.columns].copy()
        except Exception as e:
            st.sidebar.error(f"Load error: {e}")
    return blank_df()

def load_pre_soil():
    """Load pre-experiment baseline soil from saved Excel, or return empty template."""
    if os.path.exists(DATA_FILE):
        try:
            pre = pd.read_excel(DATA_FILE, sheet_name="Pre-Experiment Soil")
            for k in SOIL_KEYS:
                if k not in pre.columns:
                    pre[k] = np.nan
            if "Sample" not in pre.columns:
                pre.insert(0, "Sample", [f"C{i+1}" for i in range(len(pre))])
            return pre
        except Exception:
            pass
    return pd.DataFrame(
        [{"Sample": f"C{i+1}"} | {k: np.nan for k in SOIL_KEYS} for i in range(3)]
    )

def save_data():
    try:
        with pd.ExcelWriter(DATA_FILE, engine="openpyxl") as writer:
            st.session_state.df.to_excel(writer, sheet_name="Raw Data", index=False)
            st.session_state.pre_soil.to_excel(
                writer, sheet_name="Pre-Experiment Soil", index=False)
        return True
    except Exception as e:
        st.sidebar.error(f"Save error: {e}")
        return False

# Soil parameter keys (used for pre-experiment baseline)
SOIL_KEYS = [k for k,l,d in PARAM_GROUPS["Soil Analysis"]]

# Pre-experiment baseline soil — load from Excel or blank template
if "pre_soil" not in st.session_state:
    st.session_state.pre_soil = load_pre_soil()

if "df" not in st.session_state:
    st.session_state.df = load_data()

# Treatment descriptions — editable by user
if "t_descs" not in st.session_state:
    st.session_state.t_descs = {t["id"]: t["desc"] for t in TREATMENTS}

# Patch TREATMENTS and T_INFO so edits propagate everywhere
for _t in TREATMENTS:
    _t["desc"] = st.session_state.t_descs.get(_t["id"], _t["desc"])
    T_INFO[_t["id"]]["desc"] = _t["desc"]

def tdesc(tid):
    """Return the (possibly user-edited) description for a treatment ID."""
    return st.session_state.t_descs.get(tid, T_INFO.get(tid, {}).get("desc", tid))

# ── STATS HELPERS ─────────────────────────────────────────────────────────────

def sync_df_from_editors():
    """Pull latest values from all editor caches into the master dataframe.
    Called before any analysis so report/stats always see current data."""
    for gn, params in PARAM_GROUPS.items():
        edf_k = f"_edf_{gn}"
        if edf_k in st.session_state:
            for k, _, _ in params:
                if k in st.session_state[edf_k].columns:
                    st.session_state.df[k] = st.session_state[edf_k][k].values

def get_groups(param):
    return [
        st.session_state.df[st.session_state.df["Treatment"]==t][param]
          .dropna().astype(float).values
        for t in T_IDS
    ]

def means_se_sd(param):
    gs = get_groups(param)
    m  = np.array([np.mean(g)                        if len(g)>0 else np.nan for g in gs])
    se = np.array([np.std(g,ddof=1)/np.sqrt(len(g))  if len(g)>1 else 0.0   for g in gs])
    sd = np.array([np.std(g,ddof=1)                  if len(g)>1 else 0.0   for g in gs])
    return m, se, sd

def anova_table(param):
    gs    = get_groups(param)
    valid = [g for g in gs if len(g) >= 2]
    if len(valid) < 2:
        return None
    all_v = np.concatenate(valid)
    gm    = np.mean(all_v)
    N, k  = len(all_v), len(valid)
    ss_t  = sum(len(g)*(np.mean(g)-gm)**2   for g in valid)
    ss_e  = sum(np.sum((g-np.mean(g))**2)   for g in valid)
    df_t, df_e = k-1, N-k
    ms_t  = ss_t/df_t
    ms_e  = ss_e/df_e if df_e>0 else np.nan
    f_v   = ms_t/ms_e if ms_e and ms_e>0 else np.nan
    p_v   = 1-stats.f.cdf(f_v, df_t, df_e) if not np.isnan(f_v) else np.nan
    cv    = np.sqrt(ms_e)/gm*100 if gm!=0 and not np.isnan(ms_e) else np.nan
    # Harmonic mean of group sizes — the correct n for a single LSD when
    # replication is unequal (e.g. a pot excluded after hail damage).
    sizes  = [len(g) for g in valid]
    n_harm = len(sizes) / sum(1.0/s for s in sizes) if all(s > 0 for s in sizes) else np.nan
    return dict(ss_t=ss_t, ss_e=ss_e, ss_total=ss_t+ss_e,
                df_t=df_t, df_e=df_e, df_total=N-1,
                ms_t=ms_t, ms_e=ms_e, f=f_v, p=p_v,
                cv=cv, N=N, k=k, gm=gm, n_harm=n_harm)

def run_tukey(param):
    sub = st.session_state.df[["Treatment",param]].dropna()
    if len(sub)<4 or sub["Treatment"].nunique()<2:
        return None
    try:
        return pairwise_tukeyhsd(
            endog=sub[param].astype(float),
            groups=sub["Treatment"], alpha=0.05)
    except:
        return None

def lsd_05(an, r=None):
    """Fisher's LSD at α=0.05. Uses the harmonic mean of group sizes when
    replication is unequal (falls back to 3 only if unavailable). Note: LSD is
    NOT multiplicity-corrected — the Tukey CLD governs the letter groupings;
    LSD is reported for reference only (Piepho 2018)."""
    if not an or np.isnan(an["ms_e"]):
        return np.nan
    if r is None:
        r = an.get("n_harm", 3)
        if r is None or (isinstance(r, float) and np.isnan(r)):
            r = 3
    return stats.t.ppf(0.975, an["df_e"]) * np.sqrt(2*an["ms_e"]/r)

def sig_label(p):
    if np.isnan(p): return ""
    if p < 0.001: return "***"
    if p < 0.01:  return "**"
    if p < 0.05:  return "*"
    return "ns"

# ── COMPACT LETTER DISPLAY ────────────────────────────────────────────────────
def compact_letter_display(tukey_res, means_dict):
    """
    Correct CLD via maximal clique enumeration in the non-significance graph.
    Each maximal clique of mutually NS treatments shares a letter.
    """
    if tukey_res is None:
        return {}
    gu = list(tukey_res.groupsunique)
    n  = len(gu)

    # Build NS adjacency set
    ns = set()
    idx = 0
    for i in range(n):
        for j in range(i+1, n):
            if not tukey_res.reject[idx]:
                ns.add((gu[i], gu[j]))
                ns.add((gu[j], gu[i]))
            idx += 1

    def is_ns(a, b): return (a,b) in ns

    def is_clique(nodes):
        return all(is_ns(nodes[i], nodes[j])
                   for i in range(len(nodes))
                   for j in range(i+1, len(nodes)))

    # Enumerate all maximal cliques (n≤8, fast enough)
    maximal = []
    for size in range(n, 0, -1):
        for combo in combinations(gu, size):
            cs = set(combo)
            if is_clique(list(combo)) and not any(cs < set(mc) for mc in maximal):
                maximal.append(list(combo))

    # Sort cliques by highest mean member (descending)
    maximal.sort(key=lambda c: max(means_dict.get(g, 0) for g in c), reverse=True)

    # Assign letters
    res = {g: [] for g in gu}
    for li, clique in enumerate(maximal[:26]):
        letter = chr(ord("a") + li)
        for g in clique:
            res[g].append(letter)

    return {g: "".join(sorted(res[g])) if res[g] else "?" for g in gu}

# ── NUE ───────────────────────────────────────────────────────────────────────

# ── OUTLIER TESTS ─────────────────────────────────────────────────────────────
# Dixon's Q critical values (two-tailed, α=0.05) for n=3–10
# Dixon's Q (r10) critical values — two-sided 95% (α=0.05), Rorabacher (1991), n=3–10
DIXON_Q_CRIT = {3:0.970,4:0.829,5:0.710,6:0.625,7:0.568,8:0.526,9:0.493,10:0.466}

def dixon_q_test(vals):
    """
    Dixon's Q test (the r10 ratio) for n=3–10.
    Critical values are the two-sided 95% (α=0.05) table of Rorabacher (1991);
    the r10 statistic is used across the whole n=3–10 range, matching that table.
    Returns list of dicts for any detected outliers (low or high end).
    """
    vals = [v for v in vals if not np.isnan(v)]
    n = len(vals)
    if n < 3:
        return []
    s = sorted(vals)
    rng = s[-1] - s[0]
    if rng == 0:
        return []
    q_crit = DIXON_Q_CRIT.get(n, 0.466)
    results = []
    q_low  = (s[1]  - s[0])  / rng
    q_high = (s[-1] - s[-2]) / rng
    if q_low  > q_crit:
        results.append({"value":s[0],  "Q":round(q_low,4),
                        "Q_crit":q_crit, "end":"low",  "outlier":True})
    if q_high > q_crit:
        results.append({"value":s[-1], "Q":round(q_high,4),
                        "Q_crit":q_crit, "end":"high", "outlier":True})
    return results

def grubbs_test(vals, alpha=0.05):
    """
    Grubbs' test (ESD) for a single outlier. Returns the suspect value
    and whether it is significant at the given alpha.
    """
    vals = [v for v in vals if not np.isnan(v)]
    n = len(vals)
    if n < 3:
        return None
    arr   = np.array(vals, dtype=float)
    mean  = arr.mean()
    sd    = arr.std(ddof=1)
    if sd == 0:
        return None
    G     = np.max(np.abs(arr - mean)) / sd
    # Two-sided critical value via t-distribution
    t_crit = stats.t.ppf(1 - alpha/(2*n), df=n-2)
    G_crit = ((n-1)/np.sqrt(n)) * np.sqrt(t_crit**2/(n-2+t_crit**2))
    suspect = arr[np.argmax(np.abs(arr-mean))]
    return {"value": round(float(suspect),4),
            "G":     round(float(G),4),
            "G_crit":round(float(G_crit),4),
            "outlier": bool(G > G_crit)}

def _suspect_rep(rep_vals, suspect_val):
    """Return the replicate label (R1/R2/R3) whose value matches suspect_val."""
    for rep, v in rep_vals.items():
        if abs(float(v) - float(suspect_val)) < 1e-6:
            return rep
    return "?"

def _dixon_comment(dq_results, rep_vals, sub):
    """Human-readable explanation of a Dixon Q flag."""
    if not dq_results:
        return ""
    d     = dq_results[0]
    val   = d["value"]
    rep   = _suspect_rep(rep_vals, val)
    end   = d["end"]          # "low" or "high"
    others = [v for v in sub if abs(v - val) > 1e-6]
    mean_others = round(sum(others)/len(others), 4) if others else "—"
    pct   = abs(val - mean_others) / mean_others * 100 if mean_others != "—" and mean_others != 0 else 0
    direction = "lower" if end == "low" else "higher"
    return (
        f"{rep} ({val}) is the {end}est value — "
        f"{pct:.1f}% {direction} than the other replicates "
        f"(mean of others = {mean_others}). "
        f"Q = {d['Q']:.4f} > Q_crit {d['Q_crit']:.3f}."
    )

def _grubbs_comment(gb, rep_vals, sub):
    """Human-readable explanation of a Grubbs flag."""
    if not gb or not gb["outlier"]:
        return ""
    val   = gb["value"]
    rep   = _suspect_rep(rep_vals, val)
    mean  = round(sum(sub)/len(sub), 4)
    sd    = round(float(np.std(sub, ddof=1)), 4)
    dev   = round(abs(val - mean), 4)
    direction = "above" if val > mean else "below"
    return (
        f"{rep} ({val}) deviates most from the group mean ({mean}). "
        f"It is {dev} units ({gb['G']:.2f} SDs) {direction} the mean. "
        f"G = {gb['G']:.4f} > G_crit {gb['G_crit']:.4f}."
    )

def _clean_comment(sub, rep_vals):
    """Comment for a clean (no outlier) treatment group."""
    vals  = sorted(sub)
    rng   = round(vals[-1] - vals[0], 4)
    cv    = round(float(np.std(sub, ddof=1)) / float(np.mean(sub)) * 100, 1) if np.mean(sub) != 0 else 0
    return f"All replicates consistent. Range = {rng}, CV = {cv}%."

def outlier_scan(param):
    """
    Scan all 8 treatments for a given parameter.
    Returns a DataFrame with flags AND plain-language comments explaining
    why each value was or was not flagged.
    """
    rows = []
    for t in TREATMENTS:
        sub  = st.session_state.df[
            st.session_state.df["Treatment"]==t["id"]
        ][param].dropna().astype(float).tolist()

        if len(sub) < 3:
            rows.append({
                "Treatment": t["id"], "Description": t["desc"],
                "R1":"—","R2":"—","R3":"—",
                "Dixon Q":"n/a","Dixon Crit":"n/a","Dixon Flag":"—",
                "Grubbs G":"n/a","Grubbs Crit":"n/a","Grubbs Flag":"—",
                "Verdict":"Insufficient data",
                "Comment":"Need ≥3 replicates to run outlier tests.",
            })
            continue

        reps = st.session_state.df[
            st.session_state.df["Treatment"]==t["id"]
        ][["Replicate",param]].dropna()
        rep_vals = {r: round(v, 4) for r, v in
                    zip(reps["Replicate"].tolist(),
                        reps[param].astype(float).tolist())}

        dq = dixon_q_test(sub)
        gb = grubbs_test(sub)

        d_flag  = "⚠ OUTLIER" if dq else "✓ OK"
        g_flag  = ("⚠ OUTLIER" if gb and gb["outlier"] else "✓ OK") if gb else "—"
        flagged = bool(dq or (gb and gb["outlier"]))
        verdict = "⚠ Flag for review" if flagged else "✓ Clean"

        # ── Build comment ─────────────────────────────────────────────────────
        comment_parts = []
        if dq:
            comment_parts.append("Dixon: " + _dixon_comment(dq, rep_vals, sub))
        if gb and gb["outlier"]:
            comment_parts.append("Grubbs: " + _grubbs_comment(gb, rep_vals, sub))
        if not flagged:
            comment_parts.append(_clean_comment(sub, rep_vals))
        comment = "  |  ".join(comment_parts) if comment_parts else "—"

        rows.append({
            "Treatment":  t["id"],
            "Description":t["desc"],
            "R1": rep_vals.get("R1","—"),
            "R2": rep_vals.get("R2","—"),
            "R3": rep_vals.get("R3","—"),
            "Dixon Q":    round(dq[0]["Q"],4) if dq else "—",
            "Dixon Crit": dq[0]["Q_crit"]     if dq else DIXON_Q_CRIT.get(len(sub),0.970),
            "Dixon Flag": d_flag,
            "Grubbs G":   gb["G"]      if gb else "—",
            "Grubbs Crit":gb["G_crit"] if gb else "—",
            "Grubbs Flag":g_flag,
            "Verdict":    verdict,
            "Comment":    comment,
        })
    return pd.DataFrame(rows)

# ── ROOT:SHOOT RATIO ──────────────────────────────────────────────────────────
def compute_rs_ratio():
    """
    Compute Root:Shoot ratio per treatment.
    Uses Dry Weights (shootDW/rootDW) when available — the scientifically correct
    basis for R:S ratio. Falls back to Fresh Weights (shootWt/rootWt) if DW not yet
    entered, and notes which weight type was used.
    """
    d = st.session_state.df
    rows = []
    for t in TREATMENTS:
        sub = d[d["Treatment"]==t["id"]]
        # Prefer DW; fall back to FW
        has_dw = sub[["shootDW","rootDW"]].dropna().shape[0] > 0
        s_col  = "shootDW" if has_dw else "shootWt"
        r_col  = "rootDW"  if has_dw else "rootWt"
        wt_lbl = "DW" if has_dw else "FW*"

        s_rep = sub[["Replicate", s_col, r_col]].dropna().copy()
        if len(s_rep) > 0:
            s_rep["rs"] = s_rep[r_col] / s_rep[s_col].replace(0, np.nan)
            rs_mean = s_rep["rs"].mean()
            rs_sd   = s_rep["rs"].std(ddof=1) if len(s_rep)>1 else np.nan
            rs_se   = rs_sd/np.sqrt(len(s_rep)) if len(s_rep)>1 else np.nan
        else:
            rs_mean=rs_sd=rs_se=np.nan

        sm = sub[s_col].mean(); rm = sub[r_col].mean()
        rows.append({
            "Treatment":           t["id"],
            "Description":         t["desc"],
            f"Shoot ({wt_lbl}) g/pot": round(sm,2)     if not np.isnan(sm)     else "—",
            f"Root ({wt_lbl}) g/pot":  round(rm,2)     if not np.isnan(rm)     else "—",
            "R:S Ratio (mean)":    round(rs_mean,3)    if not np.isnan(rs_mean)else "—",
            "R:S SD":              round(rs_sd,3)      if not np.isnan(rs_sd)  else "—",
            "R:S SE":              round(rs_se,3)      if not np.isnan(rs_se)  else "—",
            "Weight basis":        "Dry weight" if has_dw else "Fresh weight (DW not yet entered)",
        })
    return pd.DataFrame(rows)

def rs_groups():
    """Return per-replicate R:S values as groups for ANOVA. Uses DW if available."""
    d = st.session_state.df
    groups = []
    for t in TREATMENTS:
        sub    = d[d["Treatment"]==t["id"]]
        has_dw = sub[["shootDW","rootDW"]].dropna().shape[0] > 0
        sc, rc = ("shootDW","rootDW") if has_dw else ("shootWt","rootWt")
        sub2   = sub[[sc, rc]].dropna()
        if len(sub2)>0:
            rs = (sub2[rc] / sub2[sc].replace(0,np.nan)).dropna().values
        else:
            rs = np.array([])
        groups.append(rs)
    return groups


# ── EFFECT SIZES ──────────────────────────────────────────────────────────────
def effect_sizes(an):
    """Compute η² and ω² from ANOVA result dict."""
    if not an or np.isnan(an.get("ss_total", np.nan)) or an["ss_total"] == 0:
        return None
    eta2  = an["ss_t"] / an["ss_total"]
    # ω² corrects for positive bias in η²; floor at 0
    omega2 = max(0.0, (an["ss_t"] - an["df_t"] * an["ms_e"])
                      / (an["ss_total"] + an["ms_e"])) if an["ms_e"] > 0 else np.nan
    def interp(v):
        if np.isnan(v): return "—"
        if v < 0.01:  return "negligible"
        if v < 0.06:  return "small"
        if v < 0.14:  return "medium"
        return "large"
    return {
        "eta2":   round(float(eta2),  4),
        "omega2": round(float(omega2),4),
        "eta2_interp":   interp(eta2),
        "omega2_interp": interp(omega2),
    }

def _es_interp(v):
    if v is None or (isinstance(v, float) and np.isnan(v)): return "—"
    if v < 0.01:  return "negligible"
    if v < 0.06:  return "small"
    if v < 0.14:  return "medium"
    return "large"

def effect_sizes_twoway(table, model=None):
    """
    Partial effect sizes for each term in a two-way ANOVA table.
      partial η²  = SS_effect / (SS_effect + SS_error)
      partial ω²  = df·(F − 1) / (df·(F − 1) + N_total)   (Olejnik & Algina 2003)
    Robust to either Type II or Type III tables and to a dropped interaction.
    """
    if table is None or "Error" not in table.index:
        return None
    ss_err = float(table.loc["Error", "sum_sq"])
    df_err = float(table.loc["Error", "df"])
    ms_err = ss_err / df_err if df_err > 0 else np.nan
    # Total sample size: prefer the fitted model's nobs (robust to Type II/III
    # table shape, which differ in whether an Intercept row is present).
    if model is not None and hasattr(model, "nobs"):
        N_total = int(model.nobs)
    else:
        N_total = int(table["df"].sum())   # fallback (Type II tables only)
    out = {}
    for src in table.index:
        if src in ("Error", "Intercept"):
            continue
        ss = float(table.loc[src, "sum_sq"])
        df = float(table.loc[src, "df"])
        F  = table.loc[src, "F"] if "F" in table.columns else np.nan
        p_eta2 = ss / (ss + ss_err) if (ss + ss_err) > 0 else np.nan
        if not np.isnan(F) and df > 0:
            num = df * (F - 1)
            p_omega2 = max(0.0, num / (num + N_total))
        else:
            p_omega2 = np.nan
        out[src] = {"p_eta2": round(float(p_eta2), 4),
                    "p_omega2": round(float(p_omega2), 4),
                    "interp": _es_interp(p_omega2)}
    return out

# ── PLANNED CONTRASTS ─────────────────────────────────────────────────────────
def planned_contrasts(param):
    """
    A-priori contrasts that answer the experiment's primary questions directly,
    using the pooled error (MS_error, df_error) from the one-way ANOVA. Pooled-
    error contrasts have more power than a full Tukey sweep and test specific
    hypotheses rather than all 28 pairwise comparisons.

    Contrasts (matched N rate, so foliar TYPE is isolated):
      • Nano vs Granular @ 50% RDN   : T7 − T8   (the headline nano-vs-gran test)
      • Nano vs Water    @ 50% RDN   : T7 − T6
      • Nano vs Water    @ 75% RDN   : T4 − T5
      • Nano vs Water    @ 100% RDN  : T3 − T2
      • Any foliar N     vs control  : (mean of T2–T8 means) − T1
    Each row: estimate, SE = sqrt(MSE·(1/n_i + 1/n_j)), t, two-sided p (df_error),
    and Holm-adjusted p across the contrast family.
    """
    an = anova_table(param)
    if an is None or np.isnan(an["ms_e"]):
        return None
    mse, df_e = an["ms_e"], an["df_e"]
    m, _, _ = means_se_sd(param)
    mean_by = {T_IDS[i]: m[i] for i in range(len(T_IDS))}
    n_by = {t: int(st.session_state.df[st.session_state.df["Treatment"]==t][param].dropna().shape[0])
            for t in T_IDS}

    def pair(a, b, label):
        ma, mb = mean_by.get(a, np.nan), mean_by.get(b, np.nan)
        na, nb = n_by.get(a, 0), n_by.get(b, 0)
        if np.isnan(ma) or np.isnan(mb) or na < 1 or nb < 1:
            return None
        est = ma - mb
        se  = np.sqrt(mse * (1.0/na + 1.0/nb))
        if se == 0:
            return None
        t   = est / se
        p   = 2 * (1 - stats.t.cdf(abs(t), df_e))
        return {"Contrast": label, "Estimate": est, "SE": se, "t": t, "p_raw": p}

    specs = [
        ("T7", "T8", "Nano vs Granular @ 50% RDN"),
        ("T7", "T6", "Nano vs Water @ 50% RDN"),
        ("T4", "T5", "Nano vs Water @ 75% RDN"),
        ("T3", "T2", "Nano vs Water @ 100% RDN"),
    ]
    rows = [r for r in (pair(a, b, lab) for a, b, lab in specs) if r is not None]

    # Fertilised (any foliar) vs absolute control
    fert = [t for t in ["T2","T3","T4","T5","T6","T7","T8"]
            if not np.isnan(mean_by.get(t, np.nan)) and n_by.get(t,0) >= 1]
    if not np.isnan(mean_by.get("T1", np.nan)) and n_by.get("T1",0) >= 1 and fert:
        mean_fert = np.mean([mean_by[t] for t in fert])
        # variance of a mean-of-means contrast vs a single group
        coef_sq = sum((1.0/len(fert))**2 / n_by[t] for t in fert) + 1.0 / n_by["T1"]
        se = np.sqrt(mse * coef_sq)
        if se > 0:
            est = mean_fert - mean_by["T1"]
            t = est / se
            rows.append({"Contrast": "Fertilised (T2–T8) vs Control (T1)",
                         "Estimate": est, "SE": se, "t": t,
                         "p_raw": 2 * (1 - stats.t.cdf(abs(t), df_e))})

    if not rows:
        return None

    # Holm step-down adjustment across the family
    order = sorted(range(len(rows)), key=lambda i: rows[i]["p_raw"])
    mtot  = len(rows)
    prev  = 0.0
    for rank, i in enumerate(order):
        adj = min(1.0, (mtot - rank) * rows[i]["p_raw"])
        adj = max(adj, prev)   # enforce monotonicity
        rows[i]["p_holm"] = adj
        prev = adj

    dec = P_DEC.get(param, 2)
    out = []
    for r in rows:
        out.append({
            "Contrast":        r["Contrast"],
            "Estimate":        round(r["Estimate"], dec),
            "SE":              round(r["SE"], dec),
            "t":               round(r["t"], 3),
            "p (raw)":         round(r["p_raw"], 4),
            "p (Holm)":        round(r["p_holm"], 4),
            "Significant":     "Yes *" if r["p_holm"] < 0.05 else "No",
        })
    return pd.DataFrame(out)

# ── BIOLOGICAL YIELD + HARVEST INDEX ─────────────────────────────────────────
def compute_bio_yield():
    """
    Biological yield = grain + straw + root dry weight (when available).
    Harvest index = grain / biological yield × 100.
    Falls back to grain + straw only when rootDW not entered.
    """
    d = st.session_state.df
    rows = []
    for t in TREATMENTS:
        sub  = d[d["Treatment"]==t["id"]]
        gy   = sub["grainY"].mean()
        sy   = sub["strawY"].mean()
        rdw  = sub["rootDW"].mean()

        has_root = not np.isnan(rdw)
        if not any(np.isnan([gy, sy, rdw])):
            bio   = gy + sy + rdw
            basis = "grain + straw + root DW"
        elif not any(np.isnan([gy, sy])):
            bio   = gy + sy
            basis = "grain + straw (root DW not entered)"
        else:
            bio   = np.nan
            basis = "—"

        hi = gy / bio * 100 if (not np.isnan(gy) and not np.isnan(bio) and bio > 0) else np.nan

        def v(x, d=2): return round(float(x), d) if not np.isnan(x) else "—"
        rows.append({
            "Treatment":              t["id"],
            "Description":           t["desc"],
            "Grain Yield (g/pot)":   v(gy),
            "Straw Yield (g/pot)":   v(sy),
            "Root DW (g/pot)":       v(rdw),
            "Biological Yield (g/pot)": v(bio),
            "Harvest Index (%)":     v(hi, 1),
            "HI Basis":              basis,
        })
    return pd.DataFrame(rows)

# ── NITROGEN BALANCE ──────────────────────────────────────────────────────────
def compute_n_balance():
    """
    Nitrogen balance per treatment (g/pot).
    N applied  = N rate (kg/ha) × pot area (m²) / 10000 × 1000
    N in crop  = grain N uptake + straw N uptake
    ΔSoil N    = post-harvest total N − pre-experiment total N  (mg/kg × 10 kg / 1000)
    N balance  = N applied − N in crop − ΔSoil N
    Positive balance = unaccounted N (losses); negative = soil N mineralisation
    """
    d          = st.session_state.df
    pot_area   = st.session_state.get("pot_area", 0.04)   # m²
    soil_wt    = 10.0                                      # kg per pot

    # Pre-experiment baseline total N
    pre  = st.session_state.pre_soil
    pre_N_mgkg = pre["totalN"].dropna().mean() if "totalN" in pre.columns else np.nan
    pre_N_gpot = pre_N_mgkg * soil_wt / 1000 if not np.isnan(pre_N_mgkg) else np.nan

    rows = []
    for t in TREATMENTS:
        sub    = d[d["Treatment"]==t["id"]]
        n_app  = t["n"] * pot_area / 10000 * 1000          # g/pot

        gy = sub["grainY"].mean(); gN = sub["grainN"].mean()
        sy = sub["strawY"].mean(); sN = sub["strawN"].mean()

        n_grain = gy*gN/100 if not any(np.isnan([gy,gN])) else np.nan
        n_straw = sy*sN/100 if not any(np.isnan([sy,sN])) else np.nan
        n_crop  = sum(x for x in [n_grain, n_straw] if not np.isnan(x))                   if not (np.isnan(n_grain) and np.isnan(n_straw)) else np.nan

        post_N_mgkg = sub["totalN"].mean()
        post_N_gpot = post_N_mgkg * soil_wt / 1000 if not np.isnan(post_N_mgkg) else np.nan
        delta_soil  = post_N_gpot - pre_N_gpot                       if not any(np.isnan([post_N_gpot, pre_N_gpot])) else np.nan

        if not np.isnan(n_crop) and not np.isnan(delta_soil):
            balance = n_app - n_crop - delta_soil
        elif not np.isnan(n_crop):
            balance = n_app - n_crop        # delta soil unknown
        else:
            balance = np.nan

        def v(x, dec=3): return round(float(x), dec) if not np.isnan(x) else "—"
        rows.append({
            "Treatment":              t["id"],
            "Description":           t["desc"],
            "N Applied (g/pot)":     v(n_app),
            "N in Grain (g/pot)":    v(n_grain),
            "N in Straw (g/pot)":    v(n_straw),
            "Total Crop N (g/pot)":  v(n_crop),
            "ΔSoil N (g/pot)":       v(delta_soil),
            "N Unaccounted (g/pot)": v(balance),
        })
    return pd.DataFrame(rows)

# ── PRE / POST SOIL COMPARISON ────────────────────────────────────────────────
def soil_comparison():
    """
    Compare pre-experiment baseline vs post-harvest treatment means
    for every soil parameter. Returns a tidy DataFrame.
    """
    d   = st.session_state.df
    pre = st.session_state.pre_soil
    soil_params = PARAM_GROUPS["Soil Analysis"]   # list of (key, label, dec)

    rows = []
    for t in TREATMENTS:
        sub = d[d["Treatment"]==t["id"]]
        row = {"Treatment": t["id"], "Description": t["desc"]}
        for k, l, dec in soil_params:
            post_m = sub[k].mean()
            pre_m  = pre[k].dropna().mean() if k in pre.columns else np.nan
            post_s = f"{post_m:.{dec}f}" if not np.isnan(post_m) else "—"
            pre_s  = f"{pre_m:.{dec}f}"  if not np.isnan(pre_m)  else "—"
            if not any(np.isnan([pre_m, post_m])) and pre_m != 0:
                pct = (post_m - pre_m) / abs(pre_m) * 100
                pct_s = f"{pct:+.1f}%"
            else:
                pct_s = "—"
            short = l.split("(")[0].strip()   # shorten label for column header
            row[f"{short} Pre"]  = pre_s
            row[f"{short} Post"] = post_s
            row[f"{short} Δ%"]   = pct_s
        rows.append(row)

    # Also return a baseline-only summary row
    baseline = {"Parameter": [], "Unit": [], "Baseline Mean": [], "SD": []}
    for k, l, dec in soil_params:
        vals = pre[k].dropna()
        baseline["Parameter"].append(l)
        baseline["Unit"].append(l.split("(")[-1].rstrip(")") if "(" in l else "")
        baseline["Baseline Mean"].append(f"{vals.mean():.{dec}f}" if len(vals)>0 else "—")
        baseline["SD"].append(f"{vals.std(ddof=1):.{dec}f}" if len(vals)>1 else "—")

    return pd.DataFrame(rows), pd.DataFrame(baseline)


# ── NUTRIENT UPTAKE (P, K, S, Ca, B) ─────────────────────────────────────────
def compute_nutrient_uptake():
    """
    Total nutrient uptake per treatment (grain + straw fractions).
    N uptake : g/pot  = yield(g) × conc(%) / 100
    P,K,S,Ca : mg/pot = yield(g) × conc(mg/kg) / 1000
    B uptake  : μg/pot = yield(g) × conc(mg/kg)
    """
    d = st.session_state.df
    NUTRIENTS = [
        ("N",  "grainN","strawN", "g/pot",  "pct"),
        ("P",  "grainP","strawP", "mg/pot", "mgkg"),
        ("K",  "grainK","strawK", "mg/pot", "mgkg"),
        ("S",  "grainS","strawS", "mg/pot", "mgkg"),
        ("Ca", "grainCa","strawCa","mg/pot","mgkg"),
        ("B",  "grainB","strawB", "μg/pot", "b"),
    ]
    rows = []
    for t in TREATMENTS:
        sub = d[d["Treatment"]==t["id"]]
        gy  = sub["grainY"].mean()
        sy  = sub["strawY"].mean()
        row = {"Treatment": t["id"], "Description": t["desc"]}
        for nut, gk, sk, unit, mode in NUTRIENTS:
            gc = sub[gk].mean(); sc = sub[sk].mean()
            if mode == "pct":
                gu = gy*gc/100        if not any(np.isnan([gy,gc])) else np.nan
                su = sy*sc/100        if not any(np.isnan([sy,sc])) else np.nan
            elif mode == "mgkg":
                gu = gy*gc/1000       if not any(np.isnan([gy,gc])) else np.nan
                su = sy*sc/1000       if not any(np.isnan([sy,sc])) else np.nan
            else:  # B in μg/pot
                gu = gy*gc            if not any(np.isnan([gy,gc])) else np.nan
                su = sy*sc            if not any(np.isnan([sy,sc])) else np.nan
            total = (gu if not np.isnan(gu) else 0) + (su if not np.isnan(su) else 0)                     if not (np.isnan(gu) and np.isnan(su)) else np.nan
            dec = 4 if mode=="pct" else 2
            def v(x): return round(float(x),dec) if not np.isnan(x) else "—"
            row[f"{nut} grain ({unit})"] = v(gu)
            row[f"{nut} straw ({unit})"] = v(su)
            row[f"{nut} total ({unit})"] = v(total)
        rows.append(row)
    return pd.DataFrame(rows)


# ── STRIP / BOX PLOT ─────────────────────────────────────────────────────────
def strip_plot(param, show_mean=True):
    """
    Individual replicate points per treatment (strip plot).
    More honest than bar+error for n=3. Mean shown as horizontal dash.
    """
    import plotly.graph_objects as go
    d     = st.session_state.df
    label = P_LABEL.get(param, param)
    dec   = P_DEC.get(param, 2)
    fig   = go.Figure()

    for i, tid in enumerate(T_IDS):
        sub  = d[d["Treatment"]==tid][param].dropna().astype(float).tolist()
        if not sub: continue
        t    = T_INFO[tid]
        col  = T_COLORS[t["type"]]
        mean = np.mean(sub)
        # Jitter x slightly so points don't overlap
        jitter = [i + (j - len(sub)/2) * 0.08 for j in range(len(sub))]
        fig.add_trace(go.Scatter(
            x=jitter, y=sub,
            mode="markers",
            name=tid,
            showlegend=False,
            marker=dict(size=12, color=col, opacity=0.8,
                        line=dict(width=1.5, color="white")),
            hovertemplate=(f"<b>{tid}</b><br>Value: %{{y:.{dec}f}}"
                           f"<extra>{t['desc']}</extra>"),
        ))
        if show_mean:
            fig.add_trace(go.Scatter(
                x=[i-0.25, i+0.25], y=[mean, mean],
                mode="lines", showlegend=False,
                line=dict(color=col, width=3),
                hoverinfo="skip",
            ))

    fig.update_layout(
        title=dict(
            text=(f"<b>{label}</b>   "
                  f"<span style='font-size:11px;color:#666'>"
                  f"Individual replicates (n=3) · dash = treatment mean</span>"),
            font=dict(size=14, family="Arial"), x=0,
        ),
        xaxis=dict(
            tickmode="array",
            tickvals=list(range(len(T_IDS))),
            ticktext=T_IDS,
            title="Treatment",
            showgrid=False,
            tickfont=dict(family="monospace", size=13),
        ),
        yaxis=dict(title=label, gridcolor="#e8e8e8"),
        plot_bgcolor="white", paper_bgcolor="white",
        height=480, margin=dict(l=70, r=60, t=80, b=60),
        font=dict(family="Arial"),
    )
    return fig


# ── SCATTER PLOT + PEARSON CORRELATION ────────────────────────────────────────
def scatter_correlation(param_x, param_y):
    """Scatter plot of two parameters with Pearson r and p-value."""
    import plotly.graph_objects as go
    d  = st.session_state.df
    lx = P_LABEL.get(param_x, param_x)
    ly = P_LABEL.get(param_y, param_y)
    dx = P_DEC.get(param_x, 2)
    dy = P_DEC.get(param_y, 2)
    fig = go.Figure()

    all_x, all_y = [], []
    for t in TREATMENTS:
        sub = d[d["Treatment"]==t["id"]]
        xs  = sub[param_x].dropna().astype(float).tolist()
        ys  = sub[param_y].dropna().astype(float).tolist()
        n   = min(len(xs), len(ys))
        if n == 0: continue
        xs, ys = xs[:n], ys[:n]
        all_x.extend(xs); all_y.extend(ys)
        col = T_COLORS[t["type"]]
        fig.add_trace(go.Scatter(
            x=xs, y=ys, mode="markers+text",
            name=t["id"],
            text=[t["id"]]*n,
            textposition="top center",
            textfont=dict(size=10, color=col),
            marker=dict(size=11, color=col, opacity=0.85,
                        line=dict(width=1.5, color="white")),
            hovertemplate=(f"<b>{t['id']}</b><br>{lx}: %{{x:.{dx}f}}"
                           f"<br>{ly}: %{{y:.{dy}f}}<extra></extra>"),
        ))

    # Pearson r
    r_val = p_val = np.nan
    if len(all_x) >= 4:
        from scipy.stats import pearsonr
        try:
            r_val, p_val = pearsonr(all_x, all_y)
            # Regression line
            m, b = np.polyfit(all_x, all_y, 1)
            x_r  = np.linspace(min(all_x), max(all_x), 50)
            fig.add_trace(go.Scatter(
                x=x_r, y=m*x_r+b,
                mode="lines", showlegend=False,
                line=dict(color="#888", width=1.5, dash="dash"),
                hoverinfo="skip",
            ))
        except: pass

    r_str = (f"r = {r_val:.3f},  p = {p_val:.4f}  {sig_label(p_val)}"
             if not np.isnan(r_val) else "")
    fig.update_layout(
        title=dict(
            text=f"<b>{lx}  vs  {ly}</b>"
                 + (f"<br><span style='font-size:11px;color:#555'>{r_str}</span>"
                    if r_str else ""),
            font=dict(size=14, family="Arial"), x=0,
        ),
        xaxis=dict(title=lx, gridcolor="#e8e8e8"),
        yaxis=dict(title=ly, gridcolor="#e8e8e8"),
        legend=dict(title="Treatment",font=dict(size=11),
                    bgcolor="rgba(255,255,255,0.9)",bordercolor="#ccc",borderwidth=1),
        plot_bgcolor="white", paper_bgcolor="white",
        height=480, margin=dict(l=70, r=60, t=90, b=60),
        font=dict(family="Arial"),
    )
    return fig, r_val, p_val


# ── DATA COMPLETENESS ────────────────────────────────────────────────────────
def data_completeness():
    """Return completeness stats per parameter group and overall."""
    d   = st.session_state.df
    out = {}
    total_cells = 0; filled_cells = 0
    for grp, params in PARAM_GROUPS.items():
        grp_total = 0; grp_filled = 0
        details = []
        for k, l, _ in params:
            if k not in d.columns: continue
            n_total  = len(d)        # 24 cells
            n_filled = d[k].notna().sum()
            grp_total  += n_total
            grp_filled += n_filled
            details.append((l, int(n_filled), n_total))
        pct = grp_filled/grp_total*100 if grp_total>0 else 0
        out[grp] = {"total":grp_total,"filled":grp_filled,"pct":pct,"details":details}
        total_cells  += grp_total
        filled_cells += grp_filled
    overall = filled_cells/total_cells*100 if total_cells>0 else 0
    return out, overall, filled_cells, total_cells

def nue_dataframe():
    """
    NUE indices on a CONSISTENT per-pot basis.

    All tissue quantities (grain/straw yield, N uptake) are entered/derived in
    g per pot. Applied N (`t["n"]`) is a field rate in kg N/ha. Mixing g/pot
    quantities with a kg/ha rate makes RE-N (a recovery fraction) and the
    PFP-N/AE-N productivity ratios dimensionally incoherent. We therefore convert
    applied N to g N per pot using the same pot-area conversion already used in
    compute_n_balance(), so numerator and denominator share one basis:

        N applied (g/pot) = rate (kg/ha) × pot_area (m²) / 10000 × 1000

    PFP-N and AE-N are then g grain per g N (= kg/kg, basis-invariant);
    RE-N is a true % recovery; PE-N (g grain gain / g N-uptake gain) and
    NHI (% of plant N in grain) do not involve applied N and were already
    basis-consistent.
    """
    d  = st.session_state.df
    pot_area = st.session_state.get("pot_area", 0.04)   # m²

    def n_app_gpot(rate_kgha):
        return rate_kgha * pot_area / 10000 * 1000      # g N / pot

    t1 = d[d["Treatment"]=="T1"]
    t1_gy = t1["grainY"].mean(); t1_sy = t1["strawY"].mean()
    t1_gN = t1["grainN"].mean(); t1_sN = t1["strawN"].mean()
    t1_nu = (t1_gy*t1_gN/100 + t1_sy*t1_sN/100
             if not any(np.isnan([t1_gy,t1_sy,t1_gN,t1_sN])) else np.nan)

    rows = []
    for t in TREATMENTS:
        sub  = d[d["Treatment"]==t["id"]]
        gy   = sub["grainY"].mean(); sy = sub["strawY"].mean()
        gN   = sub["grainN"].mean(); sN = sub["strawN"].mean()
        nu   = (gy*gN/100 + sy*sN/100
                if not any(np.isnan([gy,sy,gN,sN])) else np.nan)
        gNu  = gy*gN/100 if not any(np.isnan([gy,gN])) else np.nan
        nA_gp = n_app_gpot(t["n"]) if t["n"] > 0 else 0.0   # g N/pot

        def v(x, dec=2): return round(float(x),dec) if not np.isnan(x) else "—"

        # PFP-N, AE-N: per-pot yield (g) ÷ per-pot N (g) = g grain / g N (= kg/kg)
        pfp = v(gy/nA_gp)              if nA_gp>0 and not np.isnan(gy)              else "—"
        ae  = v((gy-t1_gy)/nA_gp)     if nA_gp>0 and not any(np.isnan([gy,t1_gy])) else "—"
        # RE-N: true recovery % — both terms now g N/pot
        re  = v((nu-t1_nu)/nA_gp*100, 1) if nA_gp>0 and not any(np.isnan([nu,t1_nu])) else "—"
        # PE-N: only interpretable when the treatment took up MORE N than control.
        dnu = nu-t1_nu if not any(np.isnan([nu,t1_nu])) else np.nan
        if (nA_gp>0 and not np.isnan(dnu) and dnu > 1e-9
                and not any(np.isnan([gy,t1_gy]))):
            pe = v((gy-t1_gy)/dnu)
        else:
            # dnu ≤ 0 → uptake not increased over control; ratio is not meaningful
            pe = "—"
        nhi = v(gNu/nu*100, 1)        if (not np.isnan(gNu) and not np.isnan(nu) and nu>0) else "—"
        pro = v(gN*5.7)               if not np.isnan(gN)                                  else "—"

        rows.append({
            "Treatment":        t["id"],
            "Description":      t["desc"],
            "N Applied (g/pot)": round(nA_gp, 4) if nA_gp>0 else "—",
            "PFP-N (g/g)":      pfp,
            "AE-N (g/g)":       ae,
            "RE-N (%)":         re,
            "PE-N (g/g)":       pe,
            "NHI (%)":          nhi,
            "Grain Protein (%)":pro,
        })
    return pd.DataFrame(rows)

# ── BAR CHART ─────────────────────────────────────────────────────────────────

# ── 2-WAY ANOVA ────────────────────────────────────────────────────────────────
def two_way_anova(param, mode="balanced"):
    """
    mode='balanced' : T2-T7 only — complete, balanced 3×2 factorial
                      (N rate {N1,N2,N3} × Foliar {Water,Nano}).
    mode='full'     : all 8 treatments. The full treatment set does NOT form a
                      complete factorial (e.g. Gran appears only at N1; Control
                      only at N0), so several N_rate × Foliar cells have zero
                      observations. With empty cells the interaction term is not
                      estimable and Type III main effects that condition on it are
                      not interpretable. In that case we DROP the interaction and
                      fit an additive main-effects model with Type II SS, which is
                      the appropriate choice for unbalanced data without a tested
                      interaction (Langsrud 2003).

    Returns (table, info). info carries: model, means_a, means_b, cell_means,
    sub, ss_type ('III' or 'II'), design_complete (bool), note (str).
    """
    d = st.session_state.df.copy()
    d["N_rate"] = d["Treatment"].map(FACTOR_N)
    d["Foliar"]  = d["Treatment"].map(FACTOR_F)

    if mode == "balanced":
        d = d[d["Treatment"].isin(BALANCED_TIDS)]

    sub = d[["N_rate","Foliar", param]].dropna()
    if len(sub) < 8 or sub["N_rate"].nunique() < 2 or sub["Foliar"].nunique() < 2:
        return None, None

    sub = sub.rename(columns={param: "Y"})
    sub["N_rate"] = sub["N_rate"].astype("category")
    sub["Foliar"]  = sub["Foliar"].astype("category")

    # ── Detect empty factorial cells ─────────────────────────────────────────
    cell_counts = sub.groupby(["N_rate", "Foliar"], observed=True)["Y"].count()
    n_levels_a  = sub["N_rate"].nunique()
    n_levels_b  = sub["Foliar"].nunique()
    n_filled    = (cell_counts > 0).sum()
    design_complete = (n_filled == n_levels_a * n_levels_b)

    try:
        if design_complete:
            # Complete factorial → test interaction with Type III SS
            model = ols("Y ~ C(N_rate) + C(Foliar) + C(N_rate):C(Foliar)",
                        data=sub).fit()
            table = anova_lm(model, typ=3)
            ss_type = "III"
            note = ("Complete factorial: interaction tested with Type III SS."
                    if mode == "balanced" else
                    "All factorial cells filled: interaction tested with Type III SS.")
        else:
            # Incomplete design → interaction not estimable. Additive model, Type II.
            model = ols("Y ~ C(N_rate) + C(Foliar)", data=sub).fit()
            table = anova_lm(model, typ=2)
            ss_type = "II"
            empty = n_levels_a * n_levels_b - n_filled
            note = (f"Design is not a complete factorial ({empty} empty cell(s)): "
                    "the N rate × Foliar interaction is not estimable, so it has "
                    "been dropped. Main effects are reported with Type II SS "
                    "(Langsrud 2003). Treat this as exploratory — for a clean "
                    "factorial use Balanced mode, and for the nano-vs-granular "
                    "comparison use the planned contrasts in the ANOVA tab.")

        # Clean up index names
        idx_map = {}
        for idx in table.index:
            if "N_rate" in str(idx) and "Foliar" not in str(idx):
                idx_map[idx] = "N Rate (Factor A)"
            elif "Foliar" in str(idx) and "N_rate" not in str(idx):
                idx_map[idx] = "Foliar Type (Factor B)"
            elif "N_rate" in str(idx) and "Foliar" in str(idx):
                idx_map[idx] = "N Rate × Foliar (Interaction)"
            elif str(idx).lower() in ("residual","error"):
                idx_map[idx] = "Error"
            elif str(idx) == "Intercept":
                idx_map[idx] = "Intercept"
            else:
                idx_map[idx] = str(idx)
        table.index = [idx_map.get(i, i) for i in table.index]

        # Factor level means
        means_a = sub.groupby("N_rate", observed=True)["Y"].agg(["mean","sem","count"]).reset_index()
        means_b = sub.groupby("Foliar", observed=True)["Y"].agg(["mean","sem","count"]).reset_index()
        cell_means = sub.groupby(["N_rate","Foliar"], observed=True)["Y"].agg(["mean","sem","count"]).reset_index()

        return table, {
            "model": model,
            "means_a": means_a,
            "means_b": means_b,
            "cell_means": cell_means,
            "sub": sub,
            "ss_type": ss_type,
            "design_complete": bool(design_complete),
            "note": note,
        }
    except Exception as e:
        return None, str(e)


def interaction_plot(param, info):
    """Plotly line chart — N Rate × Foliar interaction, publication-ready."""
    cell  = info["cell_means"].copy()
    dec   = P_DEC.get(param, 2)
    label = P_LABEL.get(param, param)
    foliar_types = sorted(cell["Foliar"].unique())
    colors_f = {"Water":"#0072B2","Nano":"#009E73","Gran":"#E69F00","Control":"#808080"}
    dashes_f  = {"Water":"solid","Nano":"dash","Gran":"dot","Control":"dashdot"}
    markers_f = {"Water":"circle","Nano":"square","Gran":"diamond","Control":"cross"}

    fig = go.Figure()
    for ft in foliar_types:
        sub = cell[cell["Foliar"]==ft].sort_values("N_rate")
        col = colors_f.get(ft,"#333")
        fig.add_trace(go.Scatter(
            x=sub["N_rate"], y=sub["mean"],
            mode="lines+markers",
            name=f"Foliar: {ft}",
            line=dict(color=col, width=2.5, dash=dashes_f.get(ft,"solid")),
            marker=dict(size=11, color=col, symbol=markers_f.get(ft,"circle"),
                        line=dict(width=1.5, color="white")),
            error_y=dict(type="data", array=sub["sem"].tolist(),
                         visible=True, color=col, thickness=1.8, width=6),
            hovertemplate=(
                f"<b>Foliar: {ft}</b><br>N Rate: %{{x}}<br>"
                f"Mean: %{{y:.{dec}f}} ± SE<extra></extra>"
            ),
        ))

    fig.update_layout(
        title=dict(
            text=(f"<b>Interaction Plot — {label}</b><br>"
                  f"<span style='font-size:11px;color:#666'>N Rate × Foliar Type  ·  Mean ± SE</span>"),
            font=dict(size=14, family="Arial, sans-serif"), x=0, xanchor="left",
        ),
        xaxis=dict(
            title=dict(text="N Rate Level", font=dict(size=13)),
            tickfont=dict(size=12),
            categoryorder="array",
            categoryarray=["N0 (0%)","N1 (50%)","N2 (75%)","N3 (100%)"],
            showgrid=False, linecolor="#aaa", linewidth=1.5, mirror=True,
        ),
        yaxis=dict(
            title=dict(text=label, font=dict(size=13)),
            tickfont=dict(size=12),
            gridcolor="#ebebeb", gridwidth=1,
            zeroline=False, linecolor="#aaa", linewidth=1.5, mirror=True,
        ),
        legend=dict(
            title=dict(text="<b>Foliar type</b>", font=dict(size=11)),
            font=dict(size=11),
            bgcolor="rgba(255,255,255,0.95)",
            bordercolor="#ccc", borderwidth=1,
            x=1.01, xanchor="left", y=1,
        ),
        plot_bgcolor="white", paper_bgcolor="white",
        height=430, margin=dict(l=75, r=200, t=85, b=70),
        font=dict(family="Arial, sans-serif"),
    )
    return fig

def df_to_jpg(df, title=""):
    """
    Render a DataFrame as a complete publication-quality Plotly table.
    Returns (fig, pixel_width, pixel_height) so the caller can pass exact
    dimensions to to_image() — preventing cut-off content.
    """
    import math
    df = df.copy().reset_index(drop=True)
    cols = list(df.columns)
    n    = len(df)

    # ── Per-column width: driven by the longest value OR header ──────────────
    CHAR_PX   = 8.5    # pixels per character at font-size 11
    HDR_PX    = 9.5    # pixels per character in bold header
    MIN_COL   = 70
    MAX_COL   = 320    # cap long-text columns so table stays printable
    PADDING   = 20     # cell left+right padding

    col_widths = []
    for col in cols:
        hdr_w  = len(str(col)) * HDR_PX + PADDING
        cell_w = max(
            (len(str(v)) * CHAR_PX + PADDING)
            for v in df[col].tolist() + [""]
        )
        col_widths.append(int(min(MAX_COL, max(MIN_COL, hdr_w, cell_w))))

    # ── Figure dimensions ─────────────────────────────────────────────────────
    # Width: sum of columns + margins + small buffer
    MARGIN_H   = 30
    fig_width  = max(820, sum(col_widths) + MARGIN_H * 2)

    # Height per row: base 32 px, but if description col exists add extra
    # to handle text wrapping in Description column
    desc_cols = [c for c in cols if "desc" in str(c).lower()]
    if desc_cols:
        # estimate wrap: description capped at MAX_COL → may wrap
        max_desc_chars = max(len(str(v)) for v in df[desc_cols[0]].tolist() + [""])
        wrap_lines = math.ceil(max_desc_chars * CHAR_PX / MAX_COL)
        row_h = max(32, wrap_lines * 16 + 12)
    else:
        row_h = 32

    HDR_H      = 40
    TITLE_H    = 52 if title else 10
    MARGIN_V   = 20
    fig_height = TITLE_H + HDR_H + n * row_h + MARGIN_V

    # ── Alternate row shading ──────────────────────────────────────────────────
    row_colors = ["#F2F7FB" if i % 2 == 0 else "#FFFFFF" for i in range(n)]

    # ── Alignment: left for text-like cols, center for numeric ────────────────
    def is_numeric_col(col):
        vals = [v for v in df[col].tolist() if str(v) not in ("—","")]
        try:
            [float(str(v).replace("*","").strip()) for v in vals[:5]]
            return True
        except:
            return False

    aligns = ["left" if (i==0 or "desc" in str(col).lower()
                          or not is_numeric_col(col))
              else "center"
              for i, col in enumerate(cols)]

    fig = go.Figure(go.Table(
        columnwidth=col_widths,
        header=dict(
            values=[f"<b>{col}</b>" for col in cols],
            fill_color="#2C4A6E",
            font=dict(color="white", size=12, family="Arial, sans-serif"),
            align="center",
            height=HDR_H,
            line=dict(color="#1a3050", width=1),
        ),
        cells=dict(
            values=[df[col].tolist() for col in cols],
            fill_color=[row_colors] * len(cols),
            font=dict(color="#1a1917", size=11, family="Arial, sans-serif"),
            align=aligns,
            height=row_h,
            line=dict(color="#d0d8e0", width=0.5),
        ),
    ))

    fig.update_layout(
        title=dict(
            text=f"<b>{title}</b>" if title else "",
            font=dict(size=12, family="Arial, sans-serif"),
            x=0, xanchor="left",
        ),
        margin=dict(l=MARGIN_H, r=MARGIN_H,
                    t=TITLE_H if title else 10, b=MARGIN_V),
        width=fig_width,
        height=fig_height,
        paper_bgcolor="white",
        autosize=False,
    )
    return fig, fig_width, fig_height


def table_jpg_btn(df, title, filename):
    """Render a download button that exports the full table as JPG."""
    try:
        fig, w, h = df_to_jpg(df, title)
        # Pass explicit pixel dimensions — prevents Plotly from guessing wrong
        jpg = fig.to_image(format="jpg", scale=2, width=w, height=h)
        st.download_button(
            "⬇ Download table as JPG",
            data=jpg,
            file_name=filename,
            mime="image/jpeg",
            key=f"jpg_{filename}_{len(df)}",
        )
    except Exception as e:
        st.caption(f"JPG export unavailable — install kaleido: `pip install kaleido`")


# ── Q-Q PLOT OF RESIDUALS ─────────────────────────────────────────────────────
def qq_plot(param):
    """
    Normal Q-Q plot of ANOVA residuals.
    Residual = observation - its treatment group mean.
    Checks normality of residuals visually — more meaningful than raw-value tests.
    """
    groups = get_groups(param)
    valid  = [(T_IDS[i], g) for i, g in enumerate(groups) if len(g) >= 1]
    if not valid:
        return None

    residuals = []
    labels    = []
    for tid, g in valid:
        gm = np.mean(g)
        for v in g:
            residuals.append(v - gm)
            labels.append(tid)

    if len(residuals) < 4:
        return None

    res_sorted = np.array(sorted(residuals))
    n          = len(res_sorted)
    probs      = [(i - 0.375) / (n + 0.25) for i in range(1, n + 1)]   # Blom formula
    theoretical = [stats.norm.ppf(p) for p in probs]

    # Reference line through 1st and 3rd quartiles (robust)
    q25_t, q75_t = np.percentile(theoretical, [25, 75])
    q25_r, q75_r = np.percentile(res_sorted,  [25, 75])
    slope     = (q75_r - q25_r) / (q75_t - q25_t) if (q75_t - q25_t) != 0 else 1
    intercept = q25_r - slope * q25_t
    x_ref     = [min(theoretical) - 0.3, max(theoretical) + 0.3]
    y_ref     = [intercept + slope * x for x in x_ref]

    # Colour points by treatment
    point_colors = [T_COLORS[T_INFO[lbl]["type"]] for lbl in labels]

    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=theoretical, y=residuals,
        mode="markers",
        marker=dict(size=10, color=point_colors,
                    line=dict(width=1, color="white")),
        text=labels,
        hovertemplate="<b>%{text}</b><br>Theoretical Q: %{x:.3f}<br>Residual: %{y:.4f}<extra></extra>",
        name="Residuals",
        showlegend=False,
    ))
    fig.add_trace(go.Scatter(
        x=x_ref, y=y_ref,
        mode="lines",
        line=dict(color="#E69F00", width=2, dash="dash"),
        name="Normal reference",
    ))

    # Shapiro-Wilk on residuals (for reference — note low n caveat in UI)
    if len(residuals) >= 3:
        sw_stat, sw_p = stats.shapiro(residuals)
        sw_note = f"Shapiro-Wilk on residuals: W={sw_stat:.4f}, p={sw_p:.4f}"
    else:
        sw_note = ""

    fig.update_layout(
        title=dict(
            text=(f"<b>Residual Q-Q Plot — {P_LABEL.get(param, param)}</b><br>"
                  f"<span style='font-size:11px;color:#666'>"
                  f"Points near the dashed line indicate approximate normality of residuals.  "
                  f"{sw_note}</span>"),
            font=dict(size=13, family="Arial"), x=0,
        ),
        xaxis=dict(title="Theoretical quantiles (Normal)",
                   tickfont=dict(size=11), showgrid=True,
                   gridcolor="#ebebeb", zeroline=True, zerolinecolor="#ccc"),
        yaxis=dict(title="Sample residuals",
                   tickfont=dict(size=11), showgrid=True, gridcolor="#ebebeb"),
        legend=dict(font=dict(size=11), bgcolor="rgba(255,255,255,0.9)",
                    bordercolor="#ccc", borderwidth=1),
        plot_bgcolor="white", paper_bgcolor="white",
        height=420, margin=dict(l=70, r=40, t=90, b=60),
        font=dict(family="Arial, sans-serif"),
    )
    return fig


# ── KRUSKAL-WALLIS + DUNN'S TEST ──────────────────────────────────────────────
def kruskal_wallis(param):
    """One-way Kruskal-Wallis H test (non-parametric ANOVA equivalent)."""
    groups = get_groups(param)
    valid  = [g for g in groups if len(g) >= 1]
    tids   = [T_IDS[i] for i, g in enumerate(groups) if len(g) >= 1]
    if len(valid) < 2:
        return None
    try:
        H, p = stats.kruskal(*valid)
        return {"H": round(float(H), 4), "p": round(float(p), 4),
                "df": len(valid) - 1, "tids": tids}
    except Exception:
        return None


def dunn_test(param):
    """
    Dunn's post-hoc pairwise test after Kruskal-Wallis.
    Uses Bonferroni correction. Manual implementation — no extra dependency.
    """
    from scipy.stats import rankdata, norm as sp_norm

    groups = get_groups(param)
    valid  = [(T_IDS[i], g) for i, g in enumerate(groups) if len(g) >= 1]
    if len(valid) < 2:
        return None

    all_vals = np.concatenate([g for _, g in valid])
    N        = len(all_vals)
    ranks    = rankdata(all_vals)

    # Mean rank per group
    ginfo = []
    idx = 0
    for tid, g in valid:
        ni = len(g)
        ginfo.append({"tid": tid, "n": ni, "R": np.mean(ranks[idx:idx + ni])})
        idx += ni

    # Tie correction factor
    _, counts = np.unique(ranks, return_counts=True)
    tie_sum   = sum(c ** 3 - c for c in counts if c > 1)
    tie_factor = (N * (N + 1) / 12) - tie_sum / (12 * (N - 1)) if N > 1 else N * (N + 1) / 12

    k             = len(ginfo)
    n_comparisons = k * (k - 1) // 2
    rows = []
    for i in range(k):
        for j in range(i + 1, k):
            gi, gj = ginfo[i], ginfo[j]
            se = np.sqrt(tie_factor * (1 / gi["n"] + 1 / gj["n"]))
            if se == 0:
                continue
            z     = abs(gi["R"] - gj["R"]) / se
            p_raw = float(2 * (1 - sp_norm.cdf(z)))
            p_bon = min(p_raw * n_comparisons, 1.0)
            rows.append({
                "Group 1":          gi["tid"],
                "Group 2":          gj["tid"],
                "Mean Rank diff":   round(abs(gi["R"] - gj["R"]), 3),
                "z":                round(z, 4),
                "p (raw)":          round(p_raw, 4),
                "p (Bonferroni)":   round(p_bon, 4),
                "Significant":      "Yes *" if p_bon < 0.05 else "No",
            })
    return pd.DataFrame(rows) if rows else None


# ── SIGNIFICANCE BRACKETS ──────────────────────────────────────────────────────
def add_significance_brackets(fig, param, mode="vs_control"):
    """
    Add significance brackets (*, **, ***) above a Plotly categorical bar chart.
    mode: 'vs_control' — only pairs vs T1
          'all'        — all significant pairs
    """
    if mode == "None":
        return fig
    tk = run_tukey(param)
    if tk is None:
        return fig

    m, se, _ = means_se_sd(param)
    gu = list(tk.groupsunique)

    # Extract significant pairs with p-values from summary table
    sig_pairs = []
    idx = 0
    try:
        rows = tk._results_table.data[1:]
        for i in range(len(gu)):
            for j in range(i + 1, len(gu)):
                if tk.reject[idx]:
                    p_adj = float(rows[idx][3])
                    sig_pairs.append((str(gu[i]), str(gu[j]), p_adj))
                idx += 1
    except Exception:
        return fig

    if mode == "vs_control":
        sig_pairs = [(a,b,p) for a,b,p in sig_pairs if "T1" in (a,b)]
    if not sig_pairs:
        return fig

    valid_tids = [T_IDS[i] for i in range(len(T_IDS)) if not np.isnan(m[i])]
    n_valid = len(valid_tids)
    if n_valid < 2:
        return fig

    bar_h  = {T_IDS[i]: (m[i] if not np.isnan(m[i]) else 0.0) for i in range(len(T_IDS))}
    bar_e  = {T_IDS[i]: se[i] for i in range(len(T_IDS))}
    max_val = max((v for v in bar_h.values() if v > 0), default=1.0)
    tick_h  = max_val * 0.03
    gap     = max_val * 0.06

    # Ceiling = top of bar + error; stacked upward for multiple brackets
    ceiling = {tid: (bar_h[tid] + bar_e[tid]) * 1.06 for tid in valid_tids}

    for a, b, p_adj in sig_pairs:
        if a not in valid_tids or b not in valid_tids:
            continue
        stars = "***" if p_adj < 0.001 else "**" if p_adj < 0.01 else "*"
        bh    = max(ceiling.get(a, 0), ceiling.get(b, 0)) + gap

        # Horizontal bar (categorical x values work with xref="x")
        fig.add_shape(type="line", x0=a, x1=b, y0=bh, y1=bh,
                      xref="x", yref="y", line=dict(color="#333", width=1.5))
        # Left tick
        fig.add_shape(type="line", x0=a, x1=a, y0=bh-tick_h, y1=bh,
                      xref="x", yref="y", line=dict(color="#333", width=1.5))
        # Right tick
        fig.add_shape(type="line", x0=b, x1=b, y0=bh-tick_h, y1=bh,
                      xref="x", yref="y", line=dict(color="#333", width=1.5))

        # Star annotation — approximate paper-x midpoint between the two bars
        ia = valid_tids.index(a)
        ib = valid_tids.index(b)
        # Bars occupy paper x ≈ 0.08 to 0.86 (empirical approximation for bargap=0.32)
        p_left, p_right = 0.08, 0.86
        step = (p_right - p_left) / max(n_valid - 1, 1)
        paper_x = p_left + (ia + ib) / 2 * step
        fig.add_annotation(xref="paper", yref="y",
                           x=paper_x, y=bh + tick_h * 0.6,
                           text=f"<b>{stars}</b>", showarrow=False,
                           font=dict(size=13, color="#333"))

        # Raise ceiling between a and b
        ia_min, ia_max = min(ia, ib), max(ia, ib)
        for tid in valid_tids[ia_min:ia_max+1]:
            ceiling[tid] = max(ceiling.get(tid, 0), bh + tick_h + gap * 0.5)

    # Extend y-axis to accommodate all brackets
    max_ceil = max(ceiling.values()) if ceiling else max_val
    fig.update_yaxes(range=[0, max_ceil * 1.15])
    return fig


# ── SIGNIFICANCE SUMMARY PROSE ────────────────────────────────────────────────
def generate_sig_summary(param):
    """
    Auto-generates a ready-to-paste results sentence pair for a parameter.
    Sentence 1: ANOVA result with effect size.
    Sentence 2: Best treatment, SD, and which treatments it outperforms.
    """
    an = anova_table(param)
    if an is None:
        return None
    m, se, sd = means_se_sd(param)
    dec   = P_DEC.get(param, 2)
    label = P_LABEL.get(param, param)
    sl    = sig_label(an["p"])
    es    = effect_sizes(an)
    lsd   = lsd_05(an)

    # Sentence 1
    es_str   = (f", ω² = {es['omega2']:.3f} [{es['omega2_interp']} effect]"
                if es else "")
    sig_word = "significantly" if an["p"] < 0.05 else "not significantly"
    s1 = (f"{label} was {sig_word} affected by treatment "
          f"(F({an['df_t']},{an['df_e']}) = {an['f']:.3f}, "
          f"p = {an['p']:.4f} {sl}{es_str}, LSD₀.₀₅ = {lsd:.{dec}f}).")

    if an["p"] >= 0.05:
        return s1

    valid = [(T_IDS[i], m[i], sd[i]) for i in range(len(T_IDS)) if not np.isnan(m[i])]
    if not valid:
        return s1

    best_tid, best_mean, best_sd = max(valid, key=lambda x: x[1])
    sd_str = f" ± {best_sd:.{dec}f} SD" if not np.isnan(best_sd) else ""
    s2     = f"The highest mean was recorded in {best_tid} ({best_mean:.{dec}f}{sd_str})"

    # Tukey comparisons involving the best treatment
    tk = run_tukey(param)
    if tk is not None:
        try:
            gu   = list(tk.groupsunique)
            rows = tk._results_table.data[1:]
            sig_vs = []
            idx = 0
            for i in range(len(gu)):
                for j in range(i + 1, len(gu)):
                    ga, gb = str(gu[i]), str(gu[j])
                    if best_tid in (ga, gb) and tk.reject[idx]:
                        p_adj = float(rows[idx][3])
                        other = gb if ga == best_tid else ga
                        sig_vs.append((other, p_adj))
                    idx += 1
            if sig_vs:
                sig_vs.sort(key=lambda x: x[1])
                pairs_str = ", ".join(f"{t} (p = {p:.3f})" for t, p in sig_vs[:6])
                if len(sig_vs) > 6:
                    pairs_str += f" and {len(sig_vs)-6} others"
                s2 += f", which was significantly superior to {pairs_str}"
        except Exception:
            pass

    s2 += "."
    return f"{s1}\n{s2}"


# ── WORD DOCUMENT REPORT ──────────────────────────────────────────────────────
def generate_docx_report():
    """
    Generate a formatted Word document (.docx) of all results.
    Returns bytes for st.download_button.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import io

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    # ── Helpers ───────────────────────────────────────────────────────────────
    def add_tbl(doc, headers, rows_data, bold_header=True):
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hdr_row = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = h
            if bold_header:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        for row_vals in rows_data:
            row = tbl.add_row()
            for i, val in enumerate(row_vals):
                row.cells[i].text = str(val)
        doc.add_paragraph()
        return tbl

    def italic_para(doc, text):
        p = doc.add_paragraph()
        p.add_run(text).italic = True
        return p

    # ── Title ─────────────────────────────────────────────────────────────────
    t = doc.add_heading("SPADE Results Report", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Statistical Platform for Agronomic Data Evaluation", 2)
    doc.add_paragraph()

    # ── Experiment details ────────────────────────────────────────────────────
    doc.add_heading("1. Experiment Details", 1)
    add_tbl(doc,
        ["Parameter", "Value"],
        [
            ["Crop variety",      "BARI Gom 33"],
            ["Experimental design","CRD — 8 Treatments × 3 Replicates (24 pots)"],
            ["Funder",            "Dhaka University Nanotechnology Centre (DUNTC)"],
            ["N rates (incl. topdress)",
             "T2/T3: 175 | T4/T5: 131.25 | T6–T8: 87.5 kg N ha⁻¹"],
            ["Basal inputs",      "TSP 150 | MOP 112 | Gypsum 125 | Boric Acid 7.5 kg ha⁻¹"],
            ["Pot soil weight",   "10 kg per pot"],
            ["Pot area",          f"{st.session_state.get('pot_area', 0.04):.3f} m²"],
        ]
    )

    # ── Treatment descriptions ─────────────────────────────────────────────────
    doc.add_heading("2. Treatment Descriptions", 1)
    add_tbl(doc,
        ["ID", "Description"],
        [[t["id"], tdesc(t["id"])] for t in TREATMENTS]
    )

    # ── Parameter results ──────────────────────────────────────────────────────
    doc.add_heading("3. Parameter Results  (Mean ± SD, n=3)", 1)
    italic_para(doc,
        "Statistical significance: * p<0.05  ** p<0.01  *** p<0.001  ns p≥0.05. "
        "Effect size: ω² negligible <0.01, small 0.01–0.06, medium 0.06–0.14, large ≥0.14. "
        "Note: p-values are per-parameter and are not corrected for testing many "
        "parameters; when screening across all parameters, treat marginal results "
        "(0.01 < p < 0.05) as exploratory or apply a false-discovery-rate adjustment."
    )

    for grp_name, params in PARAM_GROUPS.items():
        doc.add_heading(f"3.{list(PARAM_GROUPS.keys()).index(grp_name)+1}  {grp_name}", 2)
        for k, l, dec in params:
            m, se, sd = means_se_sd(k)
            has_data = any(not np.isnan(v) for v in m)
            if not has_data:
                continue
            doc.add_heading(l, 3)

            # ANOVA line
            an = anova_table(k)
            es = effect_sizes(an) if an else None
            lsd = lsd_05(an) if an else np.nan
            if an:
                sl = sig_label(an["p"])
                es_txt = (f"  ω² = {es['omega2']:.4f} [{es['omega2_interp']}]"
                          if es else "")
                p_anova = doc.add_paragraph()
                p_anova.add_run("ANOVA: ").bold = True
                p_anova.add_run(
                    f"F({an['df_t']},{an['df_e']}) = {an['f']:.3f}, "
                    f"p = {an['p']:.4f} {sl}, LSD₀.₀₅ = {lsd:.{dec}f}, "
                    f"CV = {an['cv']:.1f}%{es_txt}"
                )

            # Means table
            means_rows = []
            for i, tid in enumerate(T_IDS):
                if np.isnan(m[i]): continue
                sd_str = f"±{sd[i]:.{dec}f}" if not np.isnan(sd[i]) else "—"
                means_rows.append([tid, tdesc(tid), f"{m[i]:.{dec}f}", sd_str])
            if means_rows:
                add_tbl(doc,
                    ["Treatment", "Description", f"Mean ({l})", "SD"],
                    means_rows)

    # ── NUE indices ────────────────────────────────────────────────────────────
    doc.add_heading("4. NUE Indices", 1)
    italic_para(doc,
        "All indices on a per-pot basis. N applied converted from kg/ha to g/pot "
        "via pot area, so applied N and N uptake share one basis. "
        "PFP-N: g grain per g N applied (= kg/kg).  AE-N: g grain gain per g N applied.  "
        "RE-N: % of applied N recovered in above-ground biomass.  "
        "PE-N: g grain gain per g N-uptake gain (reported only when uptake exceeds the control).  "
        "NHI: % of plant N partitioned to grain.  Protein: grain N% × 5.7. "
        "Uptake uses oven-dry tissue mass."
    )
    nue_df = nue_dataframe()
    add_tbl(doc, list(nue_df.columns),
            [list(row) for _, row in nue_df.iterrows()])

    # ── Nutrient uptake ───────────────────────────────────────────────────────
    nu_df = compute_nutrient_uptake()
    has_nu = nu_df.drop(columns=["Treatment","Description"])                  .apply(lambda col: col.ne("—")).any().any()
    if has_nu:
        doc.add_heading("5. Multi-Nutrient Uptake", 1)
        italic_para(doc, "N in g/pot; P, K, S, Ca in mg/pot; B in μg/pot.")
        add_tbl(doc, list(nu_df.columns),
                [list(row) for _, row in nu_df.iterrows()])

    # ── N balance ─────────────────────────────────────────────────────────────
    nb_df = compute_n_balance()
    has_nb = any(nb_df["Total Crop N (g/pot)"] != "—")
    if has_nb:
        doc.add_heading("6. Nitrogen Balance", 1)
        italic_para(doc,
            "N balance = N applied − total crop N uptake − ΔSoil N. "
            "Positive = unaccounted N (losses). Negative = soil N mineralisation."
        )
        add_tbl(doc, list(nb_df.columns),
                [list(row) for _, row in nb_df.iterrows()])

    # ── Field events ──────────────────────────────────────────────────────────
    doc.add_heading("7. Field Events & Limitations", 1)
    doc.add_paragraph(
        "Hail damage: an uncontrolled hail event during grain-filling caused partial "
        "grain loss. Pots with total grain loss were excluded from yield analysis. "
        "Straw yield, tissue N, soil nutrients, and R:S ratio data remain fully valid."
    )
    doc.add_paragraph(
        "Biotic/abiotic stresses observed 8 March 2025: Fusarium head blight, "
        "frost damage, cereal leaf beetle, and aphid infestation. "
        "Recorded in experimentation diary."
    )

    # ── Statistical notes ─────────────────────────────────────────────────────
    doc.add_heading("8. Statistical Notes", 1)
    doc.add_paragraph(
        "One-way ANOVA with Tukey HSD post-hoc (Piepho 2004 maximal clique CLD, α=0.05). "
        "Two-way factorial ANOVA (N rate × foliar type) with Type III SS. "
        "Effect sizes: η² and ω² (bias-corrected). "
        "Outlier detection: Dixon's Q test + Grubbs' test. "
        "Non-parametric fallback: Kruskal-Wallis with Dunn's test (Bonferroni). "
        "All values reported as Mean ± SD (n=3 replicates per treatment). "
        "Analyses performed using SPADE — Statistical Platform for Agronomic "
        "Data Evaluation (DUNTC, 2025)."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── PARAMETER TOOLTIPS ────────────────────────────────────────────────────────
PARAM_HELP = {
    "plantH":     "Plant height from soil surface to tallest leaf/spike tip (cm). Measure 3–5 plants/pot.",
    "tillerN":    "Total tillers per pot (including main stem) at maximum tillering stage.",
    "spad_t":     "SPAD chlorophyll index on flag leaf at tillering. Higher value = more N.",
    "spad_h":     "SPAD chlorophyll index on flag leaf at heading stage.",
    "lai":        "Leaf Area Index = total leaf area / ground area at heading.",
    "stemD":      "Stem diameter 5 cm from base using vernier calipers (mm).",
    "spikeL":     "Mean spike length base-to-tip (cm). Average of 10 spikes per pot.",
    "rootL":      "Root length from crown to longest root tip (cm).",
    "shootWt":    "Fresh weight of all above-ground biomass immediately after harvest (g/pot).",
    "rootWt":     "Fresh weight of washed and blotted root system (g/pot).",
    "shootDW":    "Shoot dry weight after oven-drying at 70 °C to constant weight (g/pot).",
    "rootDW":     "Root dry weight after oven-drying at 70 °C (g/pot). Required for R:S ratio.",
    "grainsSpk":  "Mean filled grains per spike from 10 representative spikes.",
    "spikletsSpk":"Total spikelets per spike including empty ones.",
    "tgw":        "1000-Grain Weight (g): count 200 grains × 5, weigh, extrapolate.",
    "grainY":     "Oven-dried grain yield per pot (g/pot). Include shattered grain from pot surface.",
    "strawY":     "Oven-dried straw yield per pot (g/pot) = above-ground biomass minus grain.",
    "grainN":     "Grain nitrogen (%). Kjeldahl digestion. Typical wheat grain: 1.8–2.5 %.",
    "strawN":     "Straw nitrogen (%). Kjeldahl digestion. Typical: 0.4–1.0 %.",
    "grainB":     "Grain boron (mg/kg). Azomethine-H or ICP-OES after dry ashing at 550 °C.",
    "strawB":     "Straw boron (mg/kg). Same method as grain B.",
    "grainP":     "Grain phosphorus (mg/kg). Vanadomolybdate colorimetric. Typical: 3000–4500.",
    "strawP":     "Straw phosphorus (mg/kg).",
    "grainK":     "Grain potassium (mg/kg). Flame photometer after wet digest. Typical: 3500–4500.",
    "strawK":     "Straw potassium (mg/kg). Typically higher than grain K.",
    "grainS":     "Grain sulphur (mg/kg). Turbidimetry or ICP-OES. Typical: 1200–1800.",
    "strawS":     "Straw sulphur (mg/kg).",
    "grainCa":    "Grain calcium (mg/kg). AAS or ICP-OES. Typical wheat grain: 300–500.",
    "strawCa":    "Straw calcium (mg/kg). Typically much higher than grain Ca.",
    "soilPH":     "Soil pH in 1:2.5 water suspension. Optimal for wheat: 6.0–7.5.",
    "soilEC":     "Electrical conductivity (dS/m) in 1:2.5 water suspension.",
    "soilOC":     "Organic carbon (%) by Walkley-Black wet oxidation.",
    "totalN":     "Total Kjeldahl nitrogen in soil (mg/kg). Compare with pre-exp baseline.",
    "residN":     "Residual inorganic N = NH4+ + NO3- (mg/kg) by KCl extraction.",
    "availP":     "Plant-available P (mg/kg). Olsen method for neutral/alkaline soils.",
    "excK":       "Exchangeable K (cmol/kg). 1N ammonium acetate, flame photometer.",
    "availS":     "Available sulphur (mg/kg). CaCl2 extraction, BaSO4 turbidimetry.",
    "excCa":      "Exchangeable calcium (cmol/kg). Ammonium acetate, AAS or titration.",
    "hwB":        "Hot-water soluble boron (mg/kg). 100 °C water extraction, Azomethine-H.",
}

# ── DATA FLAGS ────────────────────────────────────────────────────────────────
if "data_flags" not in st.session_state:
    st.session_state.data_flags = {}

def get_flag(tid, rep, key):
    return st.session_state.data_flags.get((tid, rep, key), {"status": "", "note": ""})

def set_flag(tid, rep, key, status, note=""):
    st.session_state.data_flags[(tid, rep, key)] = {"status": status, "note": note}

def flags_dataframe():
    rows = []
    for (tid, rep, key), info in st.session_state.data_flags.items():
        if info.get("status"):
            vals = st.session_state.df[
                (st.session_state.df["Treatment"]==tid) &
                (st.session_state.df["Replicate"]==rep)
            ][key].values
            rows.append({
                "Treatment": tid, "Replicate": rep,
                "Parameter": P_LABEL.get(key, key),
                "Value": (round(float(vals[0]), P_DEC.get(key,2))
                          if len(vals)>0 and not pd.isna(vals[0]) else "—"),
                "Status": info["status"], "Note": info.get("note",""),
            })
    cols = ["Treatment","Replicate","Parameter","Value","Status","Note"]
    return pd.DataFrame(rows, columns=cols) if rows else pd.DataFrame(columns=cols)

# ── RADAR CHART ───────────────────────────────────────────────────────────────
def radar_chart(selected_params, selected_tids, normalise="max"):
    import plotly.graph_objects as go
    labels = [P_LABEL.get(k,k).split("(")[0].strip() for k in selected_params]
    if len(selected_params) < 3:
        return None
    means_all = {}
    for k in selected_params:
        mv, _, _ = means_se_sd(k)
        means_all[k] = {T_IDS[i]: mv[i] for i in range(len(T_IDS))}
    norm_means = {}
    for k in selected_params:
        vals  = [means_all[k].get(t, np.nan) for t in T_IDS]
        valid = [v for v in vals if not np.isnan(v) and v > 0]
        if not valid:
            ref = 1.0
        elif normalise == "ctrl":
            ref = means_all[k].get("T1", np.nan)
            ref = ref if (not np.isnan(ref) and ref > 0) else (max(valid) if valid else 1.0)
        else:
            ref = max(valid) if valid else 1.0
        norm_means[k] = {t: (means_all[k][t]/ref*100
                             if not np.isnan(means_all[k].get(t,np.nan)) and ref>0 else 0)
                         for t in T_IDS}
    fig = go.Figure()
    for tid in selected_tids:
        col = T_COLORS[T_INFO[tid]["type"]]
        rv  = [norm_means[k].get(tid, 0) for k in selected_params]
        rv.append(rv[0])
        fig.add_trace(go.Scatterpolar(
            r=rv, theta=labels+[labels[0]], name=tid,
            fill="toself", opacity=0.55,
            line=dict(color=col, width=2), marker=dict(size=6, color=col),
            hovertemplate="<b>%{theta}</b>: %{r:.1f}%<extra>"+tid+"</extra>",
        ))
    scale_lbl = "% of control (T1)" if normalise=="ctrl" else "% of treatment maximum"
    fig.update_layout(
        polar=dict(radialaxis=dict(visible=True, range=[0,115],
                                  ticksuffix="%", tickfont=dict(size=10)),
                   angularaxis=dict(tickfont=dict(size=11))),
        title=dict(text=(f"<b>Multi-Parameter Radar Chart</b><br>"
                         f"<span style='font-size:11px;color:#666'>"
                         f"Values normalised to {scale_lbl}</span>"),
                   font=dict(size=14,family="Arial"), x=0.5, xanchor="center"),
        legend=dict(font=dict(size=11), bgcolor="rgba(255,255,255,0.9)",
                    bordercolor="#ccc", borderwidth=1),
        height=560, margin=dict(l=60,r=60,t=90,b=60),
        paper_bgcolor="white", font=dict(family="Arial"),
    )
    return fig

def bar_chart(param, error_type="SE", letters=None):
    """Publication-ready bar chart with error bars and CLD letters."""
    m, se, sd = means_se_sd(param)
    err   = se if error_type=="SE" else sd
    label = P_LABEL.get(param, param)
    dec   = P_DEC.get(param, 2)

    fig = go.Figure()
    for i, tid in enumerate(T_IDS):
        if np.isnan(m[i]):
            continue
        t      = T_INFO[tid]
        col    = T_COLORS[t["type"]]
        col_lt = T_COLORS_LIGHT[t["type"]]
        letter = letters.get(tid, "") if letters else ""

        fig.add_trace(go.Bar(
            name=f"{tid}",
            x=[tid], y=[m[i]],
            error_y=dict(
                type="data", array=[err[i]], visible=True,
                color="#333333", thickness=2.5, width=7,
            ),
            marker=dict(
                color=col,
                line=dict(color=col, width=1.5),
                pattern=dict(shape=""),
            ),
            text=f"<b>{letter}</b>" if letter else "",
            textposition="outside",
            textfont=dict(size=15, color="#111111", family="Arial Black, Arial"),
            hovertemplate=(
                f"<b>{tid}</b><br>"
                f"{t['desc']}<br>"
                f"Mean: %{{y:.{dec}f}}<br>"
                f"±{err[i]:.{dec}f} ({error_type})"
                + (f"<br>CLD group: <b>{letter}</b>" if letter else "")
                + "<extra></extra>"
            ),
        ))

    err_label = "Standard Error (SE)" if error_type=="SE" else "Standard Deviation (SD)"
    cld_note  = "  ·  Letters = Tukey HSD grouping (α=0.05)" if letters else ""
    fig.update_layout(
        title=dict(
            text=(
                f"<b>{label}</b><br>"
                f"<span style='font-size:11px;color:#555'>"
                f"Mean ± {err_label}{cld_note}</span>"
            ),
            font=dict(size=15, family="Arial, sans-serif"),
            x=0, xanchor="left",
        ),
        xaxis=dict(
            title=dict(text="Treatment", font=dict(size=13)),
            tickfont=dict(family="Arial, sans-serif", size=13, color="#222"),
            showgrid=False,
            linecolor="#999", linewidth=1.5, mirror=True,
        ),
        yaxis=dict(
            title=dict(text=label, font=dict(size=13)),
            tickfont=dict(size=12),
            gridcolor="#e8e8e8", gridwidth=1,
            zeroline=True, zerolinecolor="#aaa", zerolinewidth=1.2,
            linecolor="#999", linewidth=1.5, mirror=True,
        ),
        showlegend=False,
        plot_bgcolor="white",
        paper_bgcolor="white",
        bargap=0.32,
        margin=dict(l=75, r=60, t=90, b=70),
        height=520,
        font=dict(family="Arial, sans-serif"),
    )

    # Add compact colour-legend as invisible scatter traces (one per type)
    _added = set()
    type_labels = {"control":"Control","water":"Water spray",
                   "nano":"Nano urea","gran":"Granulated urea spray"}
    for tid in T_IDS:
        typ = T_INFO[tid]["type"]
        if typ not in _added:
            fig.add_trace(go.Scatter(
                x=[None], y=[None], mode="markers",
                marker=dict(size=12, color=T_COLORS[typ],
                            symbol="square"),
                name=type_labels[typ],
                showlegend=True,
                legendgroup=typ,
            ))
            _added.add(typ)

    fig.update_layout(
        showlegend=True,
        legend=dict(
            title=dict(text="<b>Foliar type</b>",
                       font=dict(size=11, color="#333")),
            font=dict(size=11, family="Arial"),
            bgcolor="rgba(255,255,255,0.95)",
            bordercolor="#ccc", borderwidth=1,
            x=1.01, xanchor="left", y=1,
            orientation="v",
        ),
    )
    return fig

# ── SIDEBAR ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 🌿 SPADE")
    st.caption("Dhaka University Nanotechnology Centre")
    st.divider()

    st.markdown("**Basal fertilizers (kg/ha)**")
    st.caption("TSP: 150 | MOP: 112 | Gypsum: 125 | Boric acid: 7.5")
    st.caption("Topdress urea: 87 kg/ha (proportional to RDN)\nApplied at CRI stage — T2 to T8 only")
    st.divider()
    st.markdown("**⚠️ Field Events**")
    st.caption(
        "8 March 2025 — Multiple biotic/abiotic stresses observed:\n"
        "• Fusarium head blight\n"
        "• Frost damage\n"
        "• Cereal leaf beetle\n"
        "• Aphid infestation\n"
        "Recorded in experimentation diary."
    )
    st.divider()

    st.markdown("**Experiment settings**")
    st.number_input(
        "Pot area (m²)",
        min_value=0.005, max_value=0.5, value=0.04,
        step=0.001, format="%.3f", key="pot_area",
        help="Surface area of each pot — used for N balance (g/pot) calculation. "
             "For a 25 cm diameter pot use 0.049 m²; for 20 cm use 0.031 m²."
    )
    st.divider()
    c1, c2 = st.columns(2)
    with c1:
        if st.button("💾 Save", use_container_width=True, type="primary"):
            if save_data():
                st.success("Saved to wheat_data.xlsx")
    with c2:
        if st.button("🔄 Reload", use_container_width=True):
            st.session_state.df       = load_data()
            st.session_state.pre_soil = load_pre_soil()
            for _gn in PARAM_GROUPS:
                st.session_state.pop(f"_edf_{_gn}", None)
            st.rerun()

    st.divider()

    # Excel download (raw data + NUE)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as xw:
        st.session_state.df.to_excel(xw, sheet_name="Raw Data", index=False)
        nue_dataframe().to_excel(xw, sheet_name="NUE Indices", index=False)
        st.session_state.pre_soil.to_excel(xw, sheet_name="Pre-Experiment Soil", index=False)
        compute_bio_yield().to_excel(xw, sheet_name="Bio Yield & HI", index=False)
        compute_n_balance().to_excel(xw, sheet_name="N Balance", index=False)
        flags_dataframe().to_excel(xw, sheet_name="Data Flags", index=False)
    st.download_button(
        "⬇ Download Excel (data + NUE)",
        data=buf.getvalue(),
        file_name="wheat_data.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )
    st.divider()

    st.markdown("**Treatment colour legend**")
    for typ, col in T_COLORS.items():
        st.markdown(f"<span style='color:{col};font-size:18px'>●</span> {typ.title()}",
                    unsafe_allow_html=True)

    st.divider()
    st.markdown("**📊 Data completeness**")
    cmpl, overall_pct, filled, total = data_completeness()
    st.progress(int(overall_pct),
               text=f"Overall: {filled}/{total} cells ({overall_pct:.0f}%)")
    for grp, grp_st in cmpl.items():
        short = grp.split("&")[0].strip()
        st.caption(f"{short}: {grp_st['filled']}/{grp_st['total']} ({grp_st['pct']:.0f}%)")

# ── HEADER ────────────────────────────────────────────────────────────────────
st.markdown("## 🌿 SPADE — Statistical Platform for Agronomic Data Evaluation")
st.caption("BARI Gom 33  ·  CRD  ·  8 Treatments × 3 Replicates  ·  DUNTC Funded  |  N rates incl. topdress: T2/T3=175 · T4/T5=131.25 · T6–T8=87.5 kg N ha⁻¹")
st.divider()

# ── PARAM SELECTOR (shared) ───────────────────────────────────────────────────
param_options = {f"{l}  [{k}]": k for k,l,d in ALL_PARAMS}

# ── TABS ──────────────────────────────────────────────────────────────────────
tab_entry, tab_stats, tab_twoway, tab_nue, tab_dynamics, tab_outlier, tab_fig, tab_report = st.tabs([
    "📋 Data Entry",
    "📊 One-Way ANOVA & Tukey",
    "📊 Two-Way ANOVA",
    "⚗️ NUE & R:S Ratio",
    "🌱 Nutrient Dynamics",
    "🔍 Outlier Tests",
    "📈 Figures",
    "📄 Report Draft",
])

# ════════════════════════════════════════════════════════════════════════════════
# TAB 1 — DATA ENTRY
# ════════════════════════════════════════════════════════════════════════════════
with tab_entry:
    st.info(
        "Enter replicate values directly in the tables. "
        "**Click 💾 Save in the sidebar** after each session to write to Excel. "
        "Data reloads automatically from Excel on startup.",
        icon="ℹ️",
    )
    st.warning(
        "⚠️ Hail damage: mark grain yield pots with total loss as blank (leave empty). "
        "Straw, root, and all tissue/soil data remain fully valid.",
        icon="⚠️",
    )

    with st.expander("✏️ Edit Treatment Descriptions", expanded=False):
        st.caption(
            "Customise the treatment labels shown throughout SPADE. "
            "Changes apply immediately to all tables, figures, and the report."
        )
        desc_df = pd.DataFrame({
            "Treatment": T_IDS,
            "Description": [tdesc(tid) for tid in T_IDS],
        })
        edited_descs = st.data_editor(
            desc_df,
            column_config={
                "Treatment":   st.column_config.TextColumn("ID", disabled=True, width=70),
                "Description": st.column_config.TextColumn(
                    "Description (editable)", width="large"),
            },
            num_rows="fixed", hide_index=True,
            key="editor_t_descs",
            use_container_width=True,
        )
        for _, row in edited_descs.iterrows():
            st.session_state.t_descs[row["Treatment"]] = row["Description"]
        if st.button("Reset to defaults", key="reset_descs"):
            for t in TREATMENTS:
                st.session_state.t_descs[t["id"]] = t["desc"]
            st.rerun()

    # ── Data Quality Flags ──────────────────────────────────────────────────────
    with st.expander("🚩 Data Quality Flags & Audit Trail", expanded=False):
        st.caption("Flag individual observations for quality control. "
                   "Flags are saved to the Excel file and shown in the Outlier Tests tab.")
        _FOPTS = ["", "✓ Verified", "⚠ Suspect — review", "❌ Excluded from analysis"]
        _fq1, _fq2 = st.columns([2, 2])
        with _fq1:
            _fgrp = st.selectbox("Group:", list(PARAM_GROUPS.keys()), key="flag_grp")
        with _fq2:
            _fp_opts = {l: k for k,l,_ in PARAM_GROUPS[_fgrp]}
            _fp_lbl  = st.selectbox("Parameter:", list(_fp_opts.keys()), key="flag_param")
            _fp_key  = _fp_opts[_fp_lbl]
        _flag_rows = []
        for _ftid in T_IDS:
            for _frep in ["R1","R2","R3"]:
                _fval = st.session_state.df[
                    (st.session_state.df["Treatment"]==_ftid) &
                    (st.session_state.df["Replicate"]==_frep)
                ][_fp_key].values
                _fvs = (f"{float(_fval[0]):.{P_DEC.get(_fp_key,2)}f}"
                        if len(_fval)>0 and not pd.isna(_fval[0]) else "—")
                _cur = get_flag(_ftid, _frep, _fp_key)
                _flag_rows.append({"Trt":_ftid,"Rep":_frep,"Value":_fvs,
                                   "Status":_cur["status"],"Note":_cur.get("note","")})
        _efl = st.data_editor(
            pd.DataFrame(_flag_rows),
            column_config={
                "Trt":    st.column_config.TextColumn(disabled=True, width=55),
                "Rep":    st.column_config.TextColumn(disabled=True, width=55),
                "Value":  st.column_config.TextColumn(disabled=True, width=90),
                "Status": st.column_config.SelectboxColumn("Flag", options=_FOPTS, width=200),
                "Note":   st.column_config.TextColumn("Note", width=200),
            },
            num_rows="fixed", hide_index=True,
            key=f"flag_editor_{_fp_key}", use_container_width=True,
        )
        for _, _fr in _efl.iterrows():
            set_flag(_fr["Trt"], _fr["Rep"], _fp_key, _fr["Status"], _fr.get("Note",""))
        _afl = flags_dataframe()
        if len(_afl) > 0:
            st.markdown(f"**{len(_afl)} flag(s) currently set across all parameters:**")
            st.dataframe(_afl, hide_index=True, use_container_width=True)
            table_jpg_btn(_afl, "SPADE Data Quality Flags", "Data_Flags.jpg")

    # ── Pre-experiment baseline soil ─────────────────────────────────────────
    with st.expander("**Pre-Experiment Baseline Soil** — 3 composite samples (enter before analysis)", expanded=False):
        st.caption(
            "Enter the soil analysis values from before the experiment started. "
            "Used for pre/post comparison and nitrogen balance calculation. "
            "Three composite subsamples (C1, C2, C3) from the bulk soil used to fill pots."
        )
        soil_col_cfg = {
            "Sample": st.column_config.TextColumn("Sample", disabled=True, width=70),
        }
        for k, l, d in PARAM_GROUPS["Soil Analysis"]:
            soil_col_cfg[k] = st.column_config.NumberColumn(
                label=l, format=f"%.{d}f", min_value=0.0,
                step=round(10**(-d), d+1),
            )
        edited_pre = st.data_editor(
            st.session_state.pre_soil,
            column_config=soil_col_cfg,
            num_rows="fixed",
            use_container_width=True,
            key="editor_pre_soil",
            hide_index=True,
        )
        st.session_state.pre_soil = edited_pre.copy()

    for group_name, params in PARAM_GROUPS.items():
        keys  = [k for k,l,d in params]
        edf_k = f"_edf_{group_name}"

        # ── Key fix: initialise the editor's own dataframe in session_state
        #    once and reuse it across reruns.  Passing a freshly-built copy
        #    every rerun caused Streamlit to reset the editor before the edit
        #    was committed, forcing users to type every value twice.
        if edf_k not in st.session_state:
            st.session_state[edf_k] = (
                st.session_state.df[["Treatment","Replicate"] + keys].copy()
            )
        else:
            # Columns added in a newer version of the dashboard may be missing
            # from a cached editor that was created before the update.
            # Add any missing columns with NaN so the editor renders correctly.
            for _k in keys:
                if _k not in st.session_state[edf_k].columns:
                    st.session_state[edf_k][_k] = np.nan

        with st.expander(f"**{group_name}** — {len(params)} parameters", expanded=True):
            col_cfg = {
                "Treatment": st.column_config.TextColumn("Trt", disabled=True, width=50),
                "Replicate":  st.column_config.TextColumn("Rep", disabled=True, width=50),
            }
            for k, l, d in params:
                col_cfg[k] = st.column_config.NumberColumn(
                    label=l, format=f"%.{d}f", min_value=0.0,
                    step=round(10**(-d), d+1),
                    help=PARAM_HELP.get(k, ""),
                )

            edited = st.data_editor(
                st.session_state[edf_k],   # stable reference — never rebuilt mid-edit
                column_config=col_cfg,
                num_rows="fixed",
                use_container_width=True,
                key=f"editor_{group_name}",
                hide_index=True,
            )

            # Persist edits into both caches
            st.session_state[edf_k] = edited.copy()
            for k in keys:
                st.session_state.df[k] = edited[k].values

# ════════════════════════════════════════════════════════════════════════════════
# TAB 2 — ANOVA & TUKEY HSD
# ════════════════════════════════════════════════════════════════════════════════
with tab_stats:
    sync_df_from_editors()
    sel_label = st.selectbox("Select parameter:", list(param_options.keys()), key="stats_sel")
    param = param_options[sel_label]
    st.divider()

    an = anova_table(param)
    tk = run_tukey(param)
    m, se, sd = means_se_sd(param)

    if an is None:
        st.warning("Not enough data. Enter at least 2 replicates across 2 or more treatments.")
    else:
        # ── ANOVA summary ──
        st.markdown("#### One-Way ANOVA")
        sl = sig_label(an["p"])
        lsd = lsd_05(an)
        es  = effect_sizes(an)
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("F-statistic", f"{an['f']:.3f}")
        col2.metric("p-value",     f"{an['p']:.4f}  {sl}")
        col3.metric("CV (%)",      f"{an['cv']:.1f}" if not np.isnan(an["cv"]) else "—")
        col4.metric("LSD (0.05)",  f"{lsd:.4f}"    if not np.isnan(lsd) else "—")
        if es:
            ec1, ec2, ec3 = st.columns(3)
            ec1.metric(f"η² (eta-squared)",  f"{es['eta2']:.4f}",
                       help="Proportion of total variance explained by treatment. η² ≥ 0.01 small, ≥ 0.06 medium, ≥ 0.14 large.")
            ec2.metric(f"ω² (omega-squared)", f"{es['omega2']:.4f}",
                       help="Bias-corrected effect size. Preferred over η² for small samples.")
            ec3.metric("Effect magnitude",
                       es["omega2_interp"].title(),
                       help="Interpretation based on ω²: negligible <0.01, small 0.01–0.06, medium 0.06–0.14, large ≥0.14.")

        anova_df = pd.DataFrame({
            "Source":  ["Treatment", "Error", "Total"],
            "SS":      [f"{an['ss_t']:.4f}", f"{an['ss_e']:.4f}", f"{an['ss_total']:.4f}"],
            "df":      [an["df_t"],           an["df_e"],           an["df_total"]],
            "MS":      [f"{an['ms_t']:.4f}", f"{an['ms_e']:.4f}", ""],
            "F":       [f"{an['f']:.3f}",    "",                   ""],
            "p-value": [f"{an['p']:.4f} {sl}","",                  ""],
            "η²":      [f"{es['eta2']:.4f}" if es else "—", "", ""],
            "ω²":      [f"{es['omega2']:.4f}" if es else "—", "", ""],
        })
        st.dataframe(anova_df, hide_index=True, use_container_width=True)
        st.caption(
            "LSD (0.05) is **uncorrected** for multiple comparisons and is shown for "
            "reference only; the compact letter display below is governed by the "
            "multiplicity-controlled **Tukey HSD** (Piepho 2018). LSD uses the harmonic "
            "mean of group sizes when replication is unequal."
        )
        table_jpg_btn(
            anova_df,
            f"One-Way ANOVA — {P_LABEL.get(param, param)}  |  F({an['df_t']},{an['df_e']})={an['f']:.3f}  p={an['p']:.4f} {sl}  LSD(0.05)={lsd:.4f}  CV={an['cv']:.1f}%",
            f"ANOVA_1way_{param}.jpg",
        )
        st.divider()

        # ── Means + CLD ──
        st.markdown("#### Treatment Means ± SE  |  Compact Letter Display (Tukey HSD, α=0.05)")
        letters = {}
        if tk is not None:
            md = {T_IDS[i]: m[i] for i in range(len(T_IDS)) if not np.isnan(m[i])}
            try:
                letters = compact_letter_display(tk, md)
            except Exception as e:
                st.caption(f"CLD error: {e}")

        dec = P_DEC.get(param, 2)
        means_rows = []
        for i, tid in enumerate(T_IDS):
            means_rows.append({
                "Treatment": tid,
                "Description": T_INFO[tid]["desc"],
                f"Mean ({P_LABEL.get(param,'')})": f"{m[i]:.{dec}f}" if not np.isnan(m[i]) else "—",
                "SE":     f"±{se[i]:.{dec}f}" if not np.isnan(m[i]) else "—",
                "SD":     f"±{sd[i]:.{dec}f}" if not np.isnan(m[i]) else "—",
                "CLD":    letters.get(tid, "—"),
            })
        means_df_out = pd.DataFrame(means_rows)
        st.dataframe(means_df_out, hide_index=True, use_container_width=True)
        table_jpg_btn(means_df_out,
                      f"Treatment Means — {P_LABEL.get(param,param)}",
                      f"means_{param}.jpg")
        st.caption(
            "**CLD interpretation:** treatments sharing a letter are NOT significantly different "
            "(Tukey HSD, α=0.05). Letters computed by maximal clique method (Piepho, 2004)."
        )
        _csv_rows = []
        for _ci, _ctid in enumerate(T_IDS):
            _crep = st.session_state.df[st.session_state.df["Treatment"]==_ctid][param].dropna().tolist()
            while len(_crep) < 3: _crep.append("")
            _csv_rows.append({
                "Treatment": _ctid, "Description": tdesc(_ctid),
                "R1":_crep[0], "R2":_crep[1], "R3":_crep[2],
                "Mean": round(m[_ci],dec) if not np.isnan(m[_ci]) else "",
                "SD":   round(sd[_ci],dec) if not np.isnan(m[_ci]) else "",
                "SE":   round(se[_ci],dec) if not np.isnan(m[_ci]) else "",
                "CLD":  letters.get(_ctid,"") if letters else "",
            })
        st.download_button(
            f"⬇ {P_LABEL.get(param,param)} — per-replicate CSV",
            data=pd.DataFrame(_csv_rows).to_csv(index=False),
            file_name=f"{param}_results.csv",
            mime="text/csv", key=f"csv_{param}",
        )
        st.divider()

        # ── Tukey pairwise table ──
        if tk is not None:
            st.markdown("#### Tukey HSD Pairwise Comparisons")
            try:
                rows_tk = []
                idx = 0
                gu = tk.groupsunique
                for i in range(len(gu)):
                    for j in range(i+1, len(gu)):
                        rows_tk.append({
                            "Group 1": gu[i],
                            "Group 2": gu[j],
                            "Mean Diff": round(tk.meandiffs[idx], 4),
                            "p-adj":     round(tk.pvalues[idx] if hasattr(tk,"pvalues") else
                                               tk._results_table.data[idx+1][3], 4),
                            "Significant": "Yes ✱" if tk.reject[idx] else "No",
                        })
                        idx += 1
                tk_df = pd.DataFrame(rows_tk)
                st.dataframe(
                    tk_df.style.apply(
                        lambda col: ["color:#c0392b;font-weight:bold"
                                     if v=="Yes ✱" else "color:#999" for v in col]
                        if col.name=="Significant" else [""]*len(col),
                        axis=0),
                    hide_index=True, use_container_width=True, height=340,
                )
                table_jpg_btn(
                    tk_df,
                    f"Tukey HSD Pairwise Comparisons — {P_LABEL.get(param, param)}",
                    f"Tukey_{param}.jpg",
                )
            except Exception as e:
                st.dataframe(pd.DataFrame(tk.summary().data[1:],
                             columns=tk.summary().data[0]),
                             hide_index=True, use_container_width=True)

        # ── Planned contrasts (a-priori) ──
        st.divider()
        st.markdown("#### Planned Contrasts (a-priori, pooled error)")
        st.caption(
            "These test the experiment's specific questions directly — isolating "
            "foliar **type** at a matched N rate — using the pooled MS_error from the "
            "ANOVA above. More powerful and more targeted than the full Tukey sweep, "
            "and the cleanest way to compare nano vs granular urea (T7 vs T8 @ 50% RDN). "
            "p-values are Holm-adjusted across this contrast family."
        )
        pc_df = planned_contrasts(param)
        if pc_df is not None:
            st.dataframe(
                pc_df.style.apply(
                    lambda col: ["color:#c0392b;font-weight:bold" if v=="Yes *"
                                 else "color:#999" for v in col]
                    if col.name=="Significant" else [""]*len(col), axis=0),
                hide_index=True, use_container_width=True,
            )
            table_jpg_btn(pc_df,
                          f"Planned Contrasts — {P_LABEL.get(param, param)} (pooled error, Holm-adjusted)",
                          f"Contrasts_{param}.jpg")
        else:
            st.info("Planned contrasts need replicate data for the relevant treatments "
                    "(e.g. T7 and T8 for the nano-vs-granular comparison).")

        st.divider()
        st.markdown("#### Residual Q-Q Plot — Normality Check")
        st.caption(
            "Plots ANOVA residuals (observation − group mean) against theoretical normal quantiles. "
            "Points close to the dashed line suggest residuals are approximately normal — "
            "the key assumption for ANOVA validity. "
            "Note: with n=3 per treatment, formal normality tests have very low power; "
            "visual inspection of this plot is more informative."
        )
        fig_qq = qq_plot(param)
        if fig_qq:
            st.plotly_chart(fig_qq, use_container_width=True)
            try:
                qq_png = fig_qq.to_image(format="png", scale=2, width=900, height=450)
                st.download_button("⬇ Download Q-Q plot (PNG)", data=qq_png,
                                   file_name=f"QQ_{param}.png", mime="image/png")
            except Exception:
                st.caption("PNG export: pip install kaleido")
        else:
            st.info("Enter data for at least 4 observations to generate the Q-Q plot.")

        st.divider()
        st.markdown("#### Non-Parametric Alternative — Kruskal-Wallis H Test")
        st.caption(
            "Use this when the Q-Q plot shows clear departures from normality, "
            "or for discrete parameters (e.g. tiller count). "
            "Kruskal-Wallis makes no normality assumption — it tests whether "
            "treatment groups come from the same distribution."
        )
        kw = kruskal_wallis(param)
        if kw:
            sl_kw = sig_label(kw["p"])
            c_kw1, c_kw2, c_kw3 = st.columns(3)
            c_kw1.metric("H statistic", f"{kw['H']:.4f}")
            c_kw2.metric("df",          kw["df"])
            c_kw3.metric("p-value",     f"{kw['p']:.4f}  {sl_kw}")

            if kw["p"] < 0.05:
                st.markdown("**Significant** — running Dunn's post-hoc pairwise test "
                            "(Bonferroni correction):")
                dunn_df = dunn_test(param)
                if dunn_df is not None:
                    styled_dunn = dunn_df.style.map(
                        lambda v: "color:#c0392b;font-weight:bold" if v=="Yes *"
                                  else "color:#999",
                        subset=["Significant"]
                    )
                    st.dataframe(styled_dunn, hide_index=True,
                                 use_container_width=True, height=340)
                    table_jpg_btn(
                        dunn_df,
                        f"Dunn's Test (Bonferroni) — {P_LABEL.get(param, param)}  "
                        f"| KW H={kw['H']:.4f}  p={kw['p']:.4f} {sl_kw}",
                        f"Dunn_{param}.jpg",
                    )
            else:
                st.success("Kruskal-Wallis not significant — no post-hoc test needed.", )
        else:
            st.info("Enter data for at least 2 treatment groups to run Kruskal-Wallis.")

# ════════════════════════════════════════════════════════════════════════════════
# TAB 3 — TWO-WAY ANOVA
# ════════════════════════════════════════════════════════════════════════════════
with tab_twoway:
    sync_df_from_editors()
    st.markdown("""
**Two factors extracted from your treatment structure:**

| Factor | Levels | Treatments |
|---|---|---|
| **A — N Rate** | N1 (50%), N2 (75%), N3 (100%) | T6/T7, T4/T5, T2/T3 |
| **B — Foliar type** | Water, Nano | T2/T5/T6, T3/T4/T7 |

*The clean factorial (Balanced mode) spans N1–N3 × {Water, Nano}. The absolute
control (T1, N0) and the granulated-urea treatment (T8, Gran at N1) sit outside
this grid: including them (Full mode) makes the design non-factorial with empty
cells, so the interaction cannot be estimated — see the note below the table.
The nano-vs-granular question is answered by the planned contrasts in the ANOVA tab.*
    """)

    st.divider()
    c_p, c_m = st.columns([3, 2])
    with c_p:
        tw_label = st.selectbox("Parameter:", list(param_options.keys()), key="tw_sel")
        tw_param = param_options[tw_label]
    with c_m:
        tw_mode = st.radio(
            "Dataset:",
            ["Balanced (T2–T7 only)", "Full — all 8 (exploratory)"],
            help=(
                "Balanced uses T2–T7: a complete, equally-replicated 3×2 factorial "
                "(N rate × Foliar Water/Nano). Interaction tested with Type III SS.\n"
                "Full adds T1 (control) and T8 (Gran). This is NOT a complete "
                "factorial — empty cells make the interaction non-estimable, so SPADE "
                "drops it and reports main effects with Type II SS (exploratory only)."
            ),
        )
    mode_key = "balanced" if "Balanced" in tw_mode else "full"

    st.divider()
    tw_table, tw_info = two_way_anova(tw_param, mode=mode_key)

    if tw_table is None:
        if isinstance(tw_info, str):
            st.error(f"Analysis error: {tw_info}")
        else:
            st.warning(
                "Not enough data. Enter replicates for at least 4 treatments "
                "spanning 2 N-rate levels and 2 foliar types."
            )
    else:
        dec = P_DEC.get(tw_param, 2)
        ss_type = tw_info.get("ss_type", "III")
        note    = tw_info.get("note", "")

        # Surface the design status prominently
        if not tw_info.get("design_complete", True):
            st.warning(note, icon="⚠️")
        else:
            st.caption(note)

        # ── ANOVA table ──
        st.markdown(f"#### Two-Way ANOVA Table (Type {ss_type} SS)")
        tw_disp = tw_table.copy()
        tw_disp = tw_disp[~tw_disp.index.str.contains("Intercept", na=False)]

        rows_out = []
        for src, row in tw_disp.iterrows():
            p_val = row.get("PR(>F)", np.nan)
            sl = sig_label(p_val) if not np.isnan(p_val) else ""
            rows_out.append({
                "Source":    src,
                "SS":        f"{row.get('sum_sq', np.nan):.4f}",
                "df":        int(row.get("df", 0)),
                "MS":        f"{row.get('mean_sq', row.get('sum_sq',np.nan)/max(row.get('df',1),1)):.4f}"
                             if src != "Error" else
                             f"{row.get('sum_sq',np.nan)/max(row.get('df',1),1):.4f}",
                "F":         f"{row.get('F',np.nan):.3f}" if not np.isnan(row.get("F",np.nan)) else "—",
                "p-value":   f"{p_val:.4f} {sl}" if not np.isnan(p_val) else "—",
            })
        tw_anova_df = pd.DataFrame(rows_out)
        st.dataframe(tw_anova_df, hide_index=True, use_container_width=True)
        table_jpg_btn(
            tw_anova_df,
            f"Two-Way ANOVA (Type {ss_type} SS) — {P_LABEL.get(tw_param, tw_param)}  |  {'Balanced (T2–T7)' if mode_key=='balanced' else 'Full — all 8 (exploratory)'}",
            f"ANOVA_2way_{tw_param}.jpg",
        )

        # ── Partial effect sizes per factor ──
        es_tw = effect_sizes_twoway(tw_table, tw_info.get("model"))
        if es_tw:
            st.caption(
                "**Partial effect sizes** (partial η² = SS_effect / (SS_effect + SS_error); "
                "partial ω² from F and df, Olejnik & Algina 2003). "
                "Thresholds: small 0.01, medium 0.06, large 0.14."
            )
            es_rows = [{"Source": k,
                        "partial η²": f"{v['p_eta2']:.4f}",
                        "partial ω²": f"{v['p_omega2']:.4f}",
                        "Magnitude":  v["interp"].title()}
                       for k, v in es_tw.items()]
            st.dataframe(pd.DataFrame(es_rows), hide_index=True, use_container_width=True)

        # Quick significance summary
        def get_p(src_fragment):
            for src in tw_table.index:
                if src_fragment.lower() in src.lower():
                    return tw_table.loc[src, "PR(>F)"]
            return np.nan

        p_a   = get_p("factor a")
        p_b   = get_p("factor b")
        p_int = get_p("interaction")

        c1, c2, c3 = st.columns(3)
        c1.metric("N Rate (A)",      f"p = {p_a:.4f} {sig_label(p_a)}"   if not np.isnan(p_a)   else "—")
        c2.metric("Foliar Type (B)", f"p = {p_b:.4f} {sig_label(p_b)}"   if not np.isnan(p_b)   else "—")
        c3.metric("A × B Interaction", f"p = {p_int:.4f} {sig_label(p_int)}" if not np.isnan(p_int) else "—")

        st.divider()

        # ── Factor level means ──
        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown("#### Factor A means — N Rate")
            ma = tw_info["means_a"].copy()
            ma.columns = ["N Rate", "Mean", "SE", "n"]
            ma["Mean"] = ma["Mean"].round(dec)
            ma["SE"]   = ma["SE"].round(dec)
            ma = ma.sort_values("N Rate")
            st.dataframe(ma, hide_index=True, use_container_width=True)

        with col_b:
            st.markdown("#### Factor B means — Foliar Type")
            mb = tw_info["means_b"].copy()
            mb.columns = ["Foliar", "Mean", "SE", "n"]
            mb["Mean"] = mb["Mean"].round(dec)
            mb["SE"]   = mb["SE"].round(dec)
            st.dataframe(mb, hide_index=True, use_container_width=True)

        st.divider()

        # ── Cell means table ──
        st.markdown("#### Cell Means (N Rate × Foliar)")
        cm = tw_info["cell_means"].copy()
        cm.columns = ["N Rate", "Foliar", "Mean", "SE", "n"]
        cm["Mean"] = cm["Mean"].round(dec)
        cm["SE"]   = cm["SE"].round(dec)
        pivot = cm.pivot(index="N Rate", columns="Foliar", values="Mean")
        pivot = pivot.reindex(index=["N0 (0%)","N1 (50%)","N2 (75%)","N3 (100%)"])
        st.dataframe(pivot, use_container_width=True)
        table_jpg_btn(pivot.reset_index(),
                      f"Cell Means — {P_LABEL.get(tw_param,tw_param)} (N Rate × Foliar)",
                      f"cell_means_{tw_param}.jpg")

        st.divider()

        # ── Interaction plot ──
        st.markdown("#### Interaction Plot  (Mean ± SE)")
        fig_int = interaction_plot(tw_param, tw_info)
        st.plotly_chart(fig_int, use_container_width=True)

        # Download interaction plot
        try:
            int_png = fig_int.to_image(format="png", scale=2.5, width=1200, height=480)
            int_jpg = fig_int.to_image(format="jpg", scale=2.5, width=1200, height=480)
            ci1, ci2 = st.columns(2)
            with ci1:
                st.download_button("⬇ Interaction plot PNG",
                    data=int_png, file_name=f"{tw_param}_interaction.png",
                    mime="image/png")
            with ci2:
                st.download_button("⬇ Interaction plot JPG",
                    data=int_jpg, file_name=f"{tw_param}_interaction.jpg",
                    mime="image/jpeg")
        except Exception:
            st.caption("PNG export: pip install kaleido")

        st.divider()

        # ── Interpretation guide ──
        with st.expander("How to interpret", expanded=False):
            st.markdown("""
**Main effect of N Rate (Factor A):** Does changing the N dose (50%, 75%, 100%) significantly
affect the parameter, averaged across all foliar types?

**Main effect of Foliar Type (Factor B):** Does switching between Water, Nano, or Gran foliar
significantly affect the parameter, averaged across N rates?

**Interaction (A × B):** Does the effect of foliar type *depend on* the N rate applied,
or vice versa? A significant interaction means the two factors do not act independently —
interpret main effects cautiously and focus on cell means and the interaction plot.

**Significance codes:** *** p<0.001 · ** p<0.01 · * p<0.05 · ns p≥0.05

**Balanced mode** (recommended): Uses T2–T7, a complete 3×2 factorial (N rate ×
{Water, Nano}) with equal replication. The interaction is estimable and is tested
with Type III SS; standard factorial assumptions hold.

**Full mode** (exploratory): Adds T1 (control, 0% N) and T8 (Gran at 50% RDN). These
do not complete the factorial grid — several N rate × foliar cells are empty — so the
interaction is **not estimable**. SPADE therefore drops the interaction and reports
main effects only, with Type II SS (Langsrud 2003). Use this descriptively. The
nano-vs-granular comparison is answered cleanly by the **planned contrasts** in the
ANOVA tab (T7 vs T8 at matched 50% RDN), not by this model.
""")


# ════════════════════════════════════════════════════════════════════════════════
# TAB 3 — NUE INDICES
# ════════════════════════════════════════════════════════════════════════════════
with tab_nue:
    sync_df_from_editors()
    with st.expander("Formula reference", expanded=False):
        st.markdown("""
All indices are computed on a **consistent per-pot basis**. Applied N is converted
from the field rate (kg N/ha) to **g N/pot** using pot area, so applied N and plant
N uptake are expressed in the same units — this is required for RE-N to be a true
recovery fraction.

| Index | Formula | Unit | Requires |
|---|---|---|---|
| PFP-N | Grain yield ÷ N applied | g grain / g N (= kg/kg) | Grain yield + N rate + pot area |
| AE-N  | (Yt − YT1) ÷ N applied | g grain / g N | Grain yield + N rate + pot area |
| RE-N  | (NUt − NUT1) ÷ N applied × 100 | % | Yield + grain N% + straw N% + pot area |
| PE-N  | (Yt − YT1) ÷ (NUt − NUT1) | g/g | Same as RE |
| NHI   | Grain N uptake ÷ Total N uptake × 100 | % | Same as RE |
| Protein | Grain N% × 5.7 | % | Grain N% |

*N applied (g/pot) = rate (kg/ha) × pot area (m²) ÷ 10000 × 1000.*
*NU = total plant N uptake (g/pot) = grain DW × grain N% + straw DW × straw N%, on an oven-dry basis.*
*PE-N is shown only when a treatment's N uptake exceeds the control (NUt − NUT1 > 0); otherwise the ratio is not interpretable and is left blank.*
        """)

    nue_df = nue_dataframe()
    st.dataframe(nue_df, hide_index=True, use_container_width=True)
    table_jpg_btn(nue_df, "NUE Indices — SPADE | BARI Gom 33 Wheat Experiment", "NUE_Indices.jpg")

    buf_nue = io.BytesIO()
    nue_df.to_excel(buf_nue, index=False)
    st.download_button(
        "⬇ Download NUE table (Excel)",
        data=buf_nue.getvalue(),
        file_name="NUE_Indices.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    st.divider()
    st.markdown("#### Root : Shoot Ratio")
    st.caption(
        "Computed per-replicate then averaged. Uses **Dry Weight** (shootDW/rootDW) when available — enter oven-dried weights for the most accurate ratio. "
        "Falls back to Fresh Weight if DW not yet entered (noted in Weight basis column)."
    )
    rs_df = compute_rs_ratio()
    has_rs = any(rs_df["R:S Ratio (mean)"] != "—")
    if has_rs:
        st.dataframe(rs_df, hide_index=True, use_container_width=True)
        table_jpg_btn(rs_df,
                      "Root:Shoot Ratio — SPADE | BARI Gom 33 Wheat Experiment",
                      "RS_Ratio.jpg")

        # ANOVA on R:S ratio
        rs_grps  = rs_groups()
        valid_rs = [g for g in rs_grps if len(g)>=2]
        if len(valid_rs) >= 2:
            all_rs   = np.concatenate(valid_rs)
            gm_rs    = all_rs.mean()
            N_rs,k_rs= len(all_rs), len(valid_rs)
            ss_t_rs  = sum(len(g)*(g.mean()-gm_rs)**2 for g in valid_rs)
            ss_e_rs  = sum(np.sum((g-g.mean())**2)    for g in valid_rs)
            df_t_rs,df_e_rs = k_rs-1, N_rs-k_rs
            ms_e_rs  = ss_e_rs/df_e_rs if df_e_rs>0 else np.nan
            ms_t_rs  = ss_t_rs/df_t_rs
            f_rs     = ms_t_rs/ms_e_rs if ms_e_rs and ms_e_rs>0 else np.nan
            p_rs     = 1-stats.f.cdf(f_rs,df_t_rs,df_e_rs) if not np.isnan(f_rs) else np.nan
            cv_rs    = np.sqrt(ms_e_rs)/gm_rs*100 if gm_rs!=0 else np.nan
            lsd_rs   = stats.t.ppf(0.975,df_e_rs)*np.sqrt(2*ms_e_rs/3) if not np.isnan(ms_e_rs) else np.nan
            sl_rs    = sig_label(p_rs)
            st.markdown("**R:S Ratio — One-Way ANOVA**")
            rsa_df = pd.DataFrame({
                "Source": ["Treatment","Error","Total"],
                "SS":  [f"{ss_t_rs:.4f}",f"{ss_e_rs:.4f}",f"{ss_t_rs+ss_e_rs:.4f}"],
                "df":  [df_t_rs, df_e_rs, N_rs-1],
                "MS":  [f"{ms_t_rs:.4f}",f"{ms_e_rs:.4f}",""],
                "F":   [f"{f_rs:.3f}","",""],
                "p-value":[f"{p_rs:.4f} {sl_rs}","",""],
            })
            st.dataframe(rsa_df, hide_index=True, use_container_width=True)
            table_jpg_btn(rsa_df,
                          f"ANOVA — Root:Shoot Ratio  |  F={f_rs:.3f}  p={p_rs:.4f} {sl_rs}  "
                          f"LSD={lsd_rs:.4f}  CV={cv_rs:.1f}%",
                          "ANOVA_RS_ratio.jpg")
            c1,c2,c3,c4 = st.columns(4)
            c1.metric("F-statistic",f"{f_rs:.3f}")
            c2.metric("p-value",    f"{p_rs:.4f} {sl_rs}")
            c3.metric("LSD (0.05)", f"{lsd_rs:.4f}" if not np.isnan(lsd_rs) else "—")
            c4.metric("CV (%)",     f"{cv_rs:.1f}"  if not np.isnan(cv_rs)  else "—")
    else:
        st.info("Enter Shoot Dry Weight and Root Dry Weight in Data Entry → Growth & Biometrics.")

# ════════════════════════════════════════════════════════════════════════════════
# TAB 5 — NUTRIENT DYNAMICS
# ════════════════════════════════════════════════════════════════════════════════
with tab_dynamics:
    sync_df_from_editors()

    # ── A: Nutrient Uptake ───────────────────────────────────────────────────
    st.markdown("#### Multi-Nutrient Uptake (grain + straw fractions)")
    st.caption(
        "Uptake = yield × tissue concentration. "
        "N in g/pot; P, K, S, Ca in mg/pot; B in μg/pot. "
        "Enter tissue concentrations (grainP, strawP, etc.) in Data Entry → Harvest & Quality."
    )
    nu_df = compute_nutrient_uptake()
    has_nu = nu_df.drop(columns=["Treatment","Description"]).apply(
        lambda col: col.ne("—")).any().any()
    if has_nu:
        st.dataframe(nu_df, hide_index=True, use_container_width=True)
        table_jpg_btn(nu_df,
                      "Multi-Nutrient Uptake — SPADE | BARI Gom 33 Wheat Experiment",
                      "Nutrient_Uptake.jpg")
    else:
        st.info("Enter tissue concentrations (grainP, strawP, grainK … ) in "
                "Data Entry → Harvest & Quality to compute nutrient uptake.")

    st.divider()

    # ── B: Biological Yield & Harvest Index ──────────────────────────────────
    st.markdown("#### Biological Yield & Harvest Index")
    st.caption(
        "Biological yield = grain + straw + root dry weight (when available). "
        "Harvest Index = grain / biological yield × 100. "
        "Enter Root Dry Weight in Data Entry → Growth & Biometrics for the most accurate HI."
    )
    bio_df = compute_bio_yield()
    has_bio = any(bio_df["Biological Yield (g/pot)"] != "—")
    if has_bio:
        st.dataframe(bio_df, hide_index=True, use_container_width=True)
        table_jpg_btn(bio_df, "Biological Yield & Harvest Index — BARI Gom 33",
                      "Bio_Yield_HI.jpg")
    else:
        st.info("Enter Grain Yield and Straw Yield to compute biological yield and harvest index.")

    st.divider()

    # ── B: Nitrogen Balance ───────────────────────────────────────────────────
    st.markdown("#### Nitrogen Balance (g N per pot)")
    st.caption(
        f"N applied (g/pot) = N rate (kg/ha) × pot area ({st.session_state.get('pot_area', 0.04):.3f} m²). "
        "N balance = N applied − total crop N uptake − ΔSoil N. "
        "Positive value = unaccounted N (leaching, volatilisation, denitrification). "
        "Requires: grain N%, straw N%, post-harvest total soil N, and pre-experiment baseline soil N."
    )
    nb_df = compute_n_balance()
    has_nb = any(nb_df["Total Crop N (g/pot)"] != "—")
    if has_nb:
        st.dataframe(nb_df, hide_index=True, use_container_width=True)
        table_jpg_btn(nb_df,
                      f"N Balance (g/pot) — pot area {st.session_state.get('pot_area', 0.04):.3f} m²",
                      "N_Balance.jpg")

        # Stacked bar: N components per treatment
        nb_plot = nb_df[nb_df["N Applied (g/pot)"] != "—"].copy()
        if len(nb_plot) > 0:
            fig_nb = go.Figure()
            for col, col_color, colname in [
                ("N in Grain (g/pot)",    "#1a6b40", "N in grain"),
                ("N in Straw (g/pot)",    "#70AD47", "N in straw"),
                ("ΔSoil N (g/pot)",       "#ED7D31", "ΔSoil N"),
                ("N Unaccounted (g/pot)", "#FF0000", "Unaccounted N"),
            ]:
                vals = []
                for v in nb_plot[col].tolist():
                    try: vals.append(float(v))
                    except: vals.append(0)
                fig_nb.add_trace(go.Bar(
                    name=colname, x=nb_plot["Treatment"].tolist(),
                    y=vals, marker_color=col_color,
                ))
            fig_nb.update_layout(
                barmode="stack",
                title="N Balance Components per Treatment (g N/pot)",
                xaxis_title="Treatment", yaxis_title="g N per pot",
                plot_bgcolor="white", paper_bgcolor="white",
                legend=dict(x=1.01, xanchor="left"),
                height=420, margin=dict(l=60,r=200,t=60,b=60),
                font=dict(family="Arial"),
            )
            st.plotly_chart(fig_nb, use_container_width=False)
            try:
                nb_png = fig_nb.to_image(format="png", scale=2, width=1100, height=450)
                st.download_button("⬇ Download N balance chart (PNG)",
                                   data=nb_png, file_name="N_Balance_chart.png",
                                   mime="image/png")
            except: st.caption("PNG export: pip install kaleido")
    else:
        st.info("Enter grain N%, straw N%, and post-harvest total soil N to compute the N balance.")

    st.divider()

    # ── C: Pre / Post Soil Comparison ────────────────────────────────────────
    st.markdown("#### Pre-Experiment vs Post-Harvest Soil Comparison")
    pre = st.session_state.pre_soil
    has_pre = pre.drop(columns=["Sample"]).notna().any().any()
    if not has_pre:
        st.info(
            "Enter pre-experiment baseline soil values in **Data Entry → "
            "Pre-Experiment Baseline Soil** to enable this comparison."
        )
    else:
        comp_df, baseline_df = soil_comparison()

        st.markdown("**Baseline summary (pre-experiment, 3 composite samples)**")
        st.dataframe(baseline_df, hide_index=True, use_container_width=True)
        table_jpg_btn(baseline_df,
                      "Pre-Experiment Soil Baseline — BARI Gom 33 Experiment",
                      "Soil_Baseline.jpg")

        st.markdown("**Post-harvest vs baseline: treatment-wise change (Δ%)**")
        st.caption("Green Δ% = increase relative to baseline; negative = decrease. "
                   "Shows treatment effect on soil nutrient status.")
        st.dataframe(comp_df, hide_index=True, use_container_width=True,
                     height=min(400, (len(comp_df)+2)*38))
        table_jpg_btn(comp_df,
                      "Pre vs Post-Harvest Soil Nutrient Comparison — All Treatments",
                      "Soil_Pre_Post_Comparison.jpg")

# ════════════════════════════════════════════════════════════════════════════════
# TAB 5 — OUTLIER TESTS
# ════════════════════════════════════════════════════════════════════════════════
with tab_outlier:
    sync_df_from_editors()
    st.markdown("""
**Methods used:**
- **Dixon's Q test** (r10 ratio) — recommended for n=3–10. Tests whether the lowest or
  highest value in a group is a statistical outlier. Critical values: two-sided 95%
  table of Rorabacher (1991); Q_critical (n=3) = 0.970.
- **Grubbs' test** — tests whether the single most extreme value is an outlier.
  More powerful when n>6; included here as a secondary check.

A value is flagged if **either** test triggers. Reporting two tests with an OR rule
raises the per-group false-positive rate, so flags are a prompt to **review**, not a
rule to exclude. With n=3 replicates both tests have very low power regardless.

⚠️ A flagged value should be checked against your field diary before removal —
do not exclude automatically.
    """)

    with st.expander("Dixon Q critical values (α=0.05)", expanded=False):
        dq_ref = pd.DataFrame({
            "n": list(DIXON_Q_CRIT.keys()),
            "Q critical": list(DIXON_Q_CRIT.values()),
        })
        st.dataframe(dq_ref, hide_index=True, use_container_width=False)

    st.divider()
    ot_label = st.selectbox("Parameter to test:", list(param_options.keys()), key="ot_sel")
    ot_param = param_options[ot_label]

    scan_df = outlier_scan(ot_param)

    # Style: red background for flagged rows
    def colour_verdict(val):
        if "OUTLIER" in str(val):
            return "background-color:#fdeaea;color:#8b1a1a;font-weight:700"
        if "Flag" in str(val):
            return "background-color:#fff3cd;color:#856404;font-weight:700"
        if "OK" in str(val) or "Clean" in str(val):
            return "color:#155724"
        return ""

    styled = scan_df.style.map(
        colour_verdict,
        subset=["Dixon Flag","Grubbs Flag","Verdict"]
    )
    st.dataframe(styled, hide_index=True, use_container_width=True, height=350)
    table_jpg_btn(
        scan_df,
        f"Outlier Test Results — {P_LABEL.get(ot_param, ot_param)}  "
        f"(Dixon Q + Grubbs, α=0.05)",
        f"Outliers_{ot_param}.jpg",
    )

    n_flagged = scan_df["Verdict"].str.contains("Flag").sum()
    if n_flagged > 0:
        st.warning(
            f"**{n_flagged} treatment(s) flagged** — check your field diary for those "
            f"replicates before deciding to exclude any values.",
            icon="⚠️",
        )
    else:
        st.success("No outliers detected across all 8 treatments for this parameter.", )

    # Manual quality flags for this parameter
    _af = flags_dataframe()
    if len(_af) > 0:
        _pf = _af[_af["Parameter"] == P_LABEL.get(ot_param, ot_param)]
        if len(_pf) > 0:
            st.divider()
            st.markdown(f"**Manual quality flags — {P_LABEL.get(ot_param, ot_param)}:**")
            st.dataframe(_pf, hide_index=True, use_container_width=True)

    st.divider()
    st.markdown("#### Scan all parameters at once")
    if st.button("🔍 Run full outlier scan — all {len(ALL_PARAMS)} parameters"):
        all_flags = []
        for k,l,_ in ALL_PARAMS:
            sc = outlier_scan(k)
            flagged = sc[sc["Verdict"].str.contains("Flag")]
            for _, row in flagged.iterrows():
                all_flags.append({
                    "Parameter": l,
                    "Treatment": row["Treatment"],
                    "R1":row["R1"],"R2":row["R2"],"R3":row["R3"],
                    "Dixon Flag": row["Dixon Flag"],
                    "Grubbs Flag":row["Grubbs Flag"],
                })
        if all_flags:
            flag_df = pd.DataFrame(all_flags)
            st.warning(f"**{len(flag_df)} potential outlier(s) found across all parameters:**")
            st.dataframe(
                flag_df.style.map(colour_verdict,
                                       subset=["Dixon Flag","Grubbs Flag"]),
                hide_index=True, use_container_width=True,
            )
            table_jpg_btn(flag_df,
                          "Full Outlier Scan — All Parameters (Dixon Q + Grubbs, α=0.05)",
                          "Outliers_full_scan.jpg")
        else:
            st.success("No outliers flagged across any parameter.", )

# ════════════════════════════════════════════════════════════════════════════════
# TAB 4 — FIGURES
# ════════════════════════════════════════════════════════════════════════════════
with tab_fig:
    sync_df_from_editors()
    fig_mode = st.radio("Chart type:", ["Bar chart (mean ± SD/SE)", "Strip plot (individual replicates)", "Scatter / Correlation"], horizontal=True)
    st.divider()

    if fig_mode != "Scatter / Correlation":
        c_sel, c_opt = st.columns([3,2])
        with c_sel:
            fig_label = st.selectbox("Parameter:", list(param_options.keys()), key="fig_sel")
            fparam = param_options[fig_label]
        with c_opt:
            if fig_mode == "Bar chart (mean ± SD/SE)":
                err_type = st.radio("Error bars:", ["SE","SD"], horizontal=True)
                show_cld = st.checkbox("Show CLD letters", value=True)
                bracket_mode = st.radio(
                    "Significance brackets:",
                    ["None", "vs Control (T1)", "All significant pairs"],
                    horizontal=True,
                    help="Draws * ** *** brackets above bars for significant Tukey pairs."
                )
            else:
                err_type = "SD"; show_cld = False; bracket_mode = "None"; bracket_mode = "None"

        # Compute CLD for bar chart
        fig_letters = {}
        if show_cld:
            ftk = run_tukey(fparam)
            fm, _, _ = means_se_sd(fparam)
            if ftk is not None:
                fmd = {T_IDS[i]: fm[i] for i in range(len(T_IDS)) if not np.isnan(fm[i])}
                try: fig_letters = compact_letter_display(ftk, fmd)
                except: pass

        if fig_mode == "Bar chart (mean ± SD/SE)":
            fig = bar_chart(fparam, error_type=err_type,
                           letters=fig_letters if show_cld else None)
            if bracket_mode != "None":
                bk_mode = "vs_control" if "Control" in bracket_mode else "all"
                fig = add_significance_brackets(fig, fparam, bk_mode)
        else:
            fig = strip_plot(fparam)

        st.plotly_chart(fig, use_container_width=False)

        try:
            img_bytes_png = fig.to_image(format="png", scale=2.5, width=1500, height=580)
            img_bytes_jpg = fig.to_image(format="jpg", scale=2.5, width=1500, height=580)
            col_png, col_jpg = st.columns(2)
            with col_png:
                st.download_button("⬇ PNG (high-res)", data=img_bytes_png,
                    file_name=f"{fparam}_{fig_mode[:3]}.png", mime="image/png")
            with col_jpg:
                st.download_button("⬇ JPG (print-ready)", data=img_bytes_jpg,
                    file_name=f"{fparam}_{fig_mode[:3]}.jpg", mime="image/jpeg")
        except Exception:
            st.caption("PNG/JPG export: pip install kaleido")

        fan = anova_table(fparam)
        if fan:
            es_fig = effect_sizes(fan)
            sl = sig_label(fan["p"]); flsd = lsd_05(fan)
            eta_str = f" | η²={es_fig['eta2']:.4f} ω²={es_fig['omega2']:.4f}" if es_fig else ""
            st.caption(
                f"**Quick stats:** F={fan['f']:.3f}, p={fan['p']:.4f} {sl} | "
                f"LSD={flsd:.4f} | CV={fan['cv']:.1f}%{eta_str}"
                if not np.isnan(fan["cv"]) else
                f"**Quick stats:** F={fan['f']:.3f}, p={fan['p']:.4f} {sl}{eta_str}"
            )

    elif fig_mode.startswith("Scatter"):  # Scatter / Correlation
        sc1, sc2 = st.columns(2)
        with sc1:
            sx_label = st.selectbox("X axis:", list(param_options.keys()), key="sc_x")
            sx_param = param_options[sx_label]
        with sc2:
            sy_label = st.selectbox("Y axis:", list(param_options.keys()),
                                   index=1, key="sc_y")
            sy_param = param_options[sy_label]
        fig_sc, r_val, p_val = scatter_correlation(sx_param, sy_param)
        st.plotly_chart(fig_sc, use_container_width=True)
        if not np.isnan(r_val):
            sl_sc = sig_label(p_val)
            strength = ("strong" if abs(r_val)>=0.7 else
                        "moderate" if abs(r_val)>=0.4 else "weak")
            direction = "positive" if r_val>0 else "negative"
            st.caption(
                f"Pearson r = {r_val:.4f}, p = {p_val:.4f} {sl_sc} · "
                f"{strength.title()} {direction} relationship · n = 24 observations"
            )
        try:
            sc_png = fig_sc.to_image(format="png", scale=2, width=900, height=520)
            st.download_button("⬇ Download scatter plot (PNG)", data=sc_png,
                file_name=f"scatter_{sx_param}_vs_{sy_param}.png", mime="image/png")
        except: st.caption("PNG export: pip install kaleido")

    else:  # 🕸 Radar (multi-parameter)
        st.info("Select ≥ 3 parameters and ≥ 2 treatments. Values are normalised to a 0–100 % scale so different units can be compared.", icon="ℹ️")
        _rdc1, _rdc2, _rdc3 = st.columns([3, 2, 2])
        with _rdc1:
            _def_keys = [k for k in ["grainY","strawY","tgw","grainN","plantH","tillerN"] if k in P_KEYS]
            _def_lbls = [P_LABEL[k] for k in _def_keys]
            _all_lbls = [l for _k,l,_ in ALL_PARAMS]
            _rd_lbls  = st.multiselect("Parameters (≥3):", _all_lbls, default=_def_lbls, key="radar_params")
        with _rdc2:
            _rd_tids  = st.multiselect("Treatments (≥2):", T_IDS, default=T_IDS, key="radar_tids")
        with _rdc3:
            _rd_norm  = st.radio("Scale:", ["% of max treatment","% of T1 (control)"], key="radar_norm")
        _lbl_to_key  = {l: k for k,l,_ in ALL_PARAMS}
        _sel_params  = [_lbl_to_key[l] for l in _rd_lbls if l in _lbl_to_key]
        _norm_mode   = "ctrl" if "T1" in _rd_norm else "max"
        if len(_sel_params) >= 3 and len(_rd_tids) >= 2:
            fig_rd = radar_chart(_sel_params, _rd_tids, normalise=_norm_mode)
            if fig_rd:
                st.plotly_chart(fig_rd, use_container_width=True)
                try:
                    _rdp = fig_rd.to_image(format="png", scale=2, width=720, height=600)
                    _rdj = fig_rd.to_image(format="jpg", scale=2, width=720, height=600)
                    _c1, _c2 = st.columns(2)
                    with _c1:
                        st.download_button("⬇ PNG", data=_rdp, file_name="radar_chart.png",
                                           mime="image/png", key="rd_png")
                    with _c2:
                        st.download_button("⬇ JPG", data=_rdj, file_name="radar_chart.jpg",
                                           mime="image/jpeg", key="rd_jpg")
                except:
                    st.caption("Export: pip install kaleido")
        else:
            st.warning("Select at least 3 parameters and 2 treatments to render the radar chart.")

# ════════════════════════════════════════════════════════════════════════════════
# TAB 5 — REPORT DRAFT
# ════════════════════════════════════════════════════════════════════════════════
with tab_report:
    sync_df_from_editors()
    rc1, rc2, rc3 = st.columns([2,2,3])
    with rc1:
        if st.button("🔄 Generate / Refresh report"):
            st.session_state.pop("report_text", None)
            st.session_state.report_ready = True
    with rc2:
        try:
            docx_bytes = generate_docx_report()
            st.download_button(
                "⬇ Download .docx report",
                data=docx_bytes,
                file_name="SPADE_Results_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as _de:
            st.caption(f"docx error: {_de}")
    with rc3:
        _sig_param_label = st.selectbox(
            "Significance summary for:",
            list(param_options.keys()), key="sig_sum_sel"
        )

    if st.session_state.get("report_ready", False):
        lines = []
        bar  = "═"*68
        dbar = "─"*68
        sbar = "·"*68

        lines += [bar,
                  "RESULTS — NANO-UREA WHEAT EXPERIMENT (BARI Gom 33)",
                  "SPADE · CRD · 8 Treatments × 3 Replicates · DUNTC Funded Project",
                  bar, ""]

        lines += ["BASAL FERTILIZER RATES & N APPLICATION", dbar,
                  "TSP 150 kg/ha | MOP 112 kg/ha | Gypsum 125 kg/ha | Boric Acid 7.5 kg/ha",
                  "Topdress urea: 87 kg/ha applied proportionally at CRI stage (T2–T8)",
                  "Total N applied per treatment (RDN + proportional topdress):",
                  "  T1: 0 kg N/ha (absolute control — no topdress)",
                  "  T2, T3: 175.00 kg N/ha  (130 + 45)",
                  "  T4, T5: 131.25 kg N/ha  (97.5 + 33.75)",
                  "  T6, T7, T8: 87.50 kg N/ha  (65 + 22.5)", ""]

        lines += ["TREATMENT DESCRIPTIONS", dbar]
        for t in TREATMENTS:
            lines.append(f"  {t['id']}:  {t['desc']}")
        lines.append("")

        # ── Loop through ALL parameter groups, show every param with data ──────
        for group_name, params in PARAM_GROUPS.items():
            group_header_written = False

            for kp, kl, _ in params:
                m2, se2, sd2 = means_se_sd(kp)
                has_data = any(not np.isnan(v) for v in m2)
                if not has_data:
                    continue

                # Write group header once
                if not group_header_written:
                    lines += ["", f"{'━'*4}  {group_name.upper()}  {'━'*(60-len(group_name))}"]
                    group_header_written = True

                dec = P_DEC.get(kp, 2)
                an  = anova_table(kp)

                if an is not None:
                    sl  = sig_label(an["p"])
                    lsd = lsd_05(an)
                    lines += [
                        "",
                        f"{kl}",
                        dbar,
                        f"  ANOVA:  F({an['df_t']},{an['df_e']}) = {an['f']:.3f}  "
                        f"p = {an['p']:.4f} {sl}  "
                        f"LSD(0.05) = {lsd:.4f}  CV = {an['cv']:.1f}%",
                        f"  Mean ± SD  (n=3 replicates per treatment)",
                    ]
                else:
                    lines += ["", f"{kl}", dbar,
                               "  ANOVA: insufficient data (need ≥2 reps per treatment)"]

                # Treatment means
                for i, tid in enumerate(T_IDS):
                    if not np.isnan(m2[i]):
                        n_reps = sum(1 for v in
                                     st.session_state.df[
                                         st.session_state.df["Treatment"]==tid][kp]
                                     .dropna())
                        sd_str = f" ± {sd2[i]:.{dec}f} SD" if n_reps > 1 else " (n=1)"
                        lines.append(f"    {tid}:  {m2[i]:.{dec}f}{sd_str}")

        lines.append("")

        # ── NUE indices ──────────────────────────────────────────────────────
        nue = nue_dataframe()
        nue_has = nue.iloc[:, 3:].apply(lambda col: col != "—").any(axis=None)
        if nue_has:
            lines += ["", f"{'━'*4}  NUE INDICES  {'━'*51}", ""]
            header = f"  {'Trt':<5} {'PFP-N':>8} {'AE-N':>8} {'RE-N%':>8} {'PE-N':>8} {'NHI%':>8} {'Protein%':>10}"
            lines.append(header)
            lines.append("  " + "─"*58)
            for _, row in nue.iterrows():
                lines.append(
                    f"  {row['Treatment']:<5}"
                    f" {str(row['PFP-N (g/g)']):>8}"
                    f" {str(row['AE-N (g/g)']):>8}"
                    f" {str(row['RE-N (%)']):>8}"
                    f" {str(row['PE-N (g/g)']):>8}"
                    f" {str(row['NHI (%)']):>8}"
                    f" {str(row['Grain Protein (%)']):>10}"
                )
            lines.append("")

        # ── Field events note ────────────────────────────────────────────────
        lines += [f"{'━'*4}  FIELD EVENTS & LIMITATIONS  {'━'*36}", "",
                  "  [1] BIOTIC/ABIOTIC STRESS — 8 March 2025",
                  "      Multiple stresses observed and recorded in experimentation diary:",
                  "      • Fusarium head blight (fungal disease)",
                  "      • Frost damage (abiotic)",
                  "      • Cereal leaf beetle (insect pest)",
                  "      • Aphid infestation (insect pest)",
                  "      These were noted as uncontrolled variables during grain-filling.",
                  "      Potential effects on yield and grain quality should be acknowledged",
                  "      in the discussion section.", "",
                  "  [2] HAIL DAMAGE",
                  "      An uncontrolled hail event during grain-filling caused partial",
                  "      grain loss. Pots with total grain loss were excluded from yield",
                  "      analysis. Straw yield, tissue N, soil nutrients, and B data",
                  "      remain fully valid and form the primary analytical contribution.",
                  ""]

        # ── Parameter coverage summary ───────────────────────────────────────
        total_params  = len(ALL_PARAMS)
        filled_params = sum(1 for k,_,_ in ALL_PARAMS
                            if any(not np.isnan(v) for v in means_se_sd(k)[0]))
        lines += [bar,
                  f"  Parameters with data: {filled_params} / {total_params}",
                  f"  Generated by SPADE (Statistical Platform for Agronomic Data Evaluation) — DUNTC",
                  bar]

        report_text = "\n".join(lines)
        st.session_state.report_text = report_text
        st.text_area("Report preview", report_text, height=560,
                     key="report_area")
        st.caption(f"Parameters with data: **{filled_params}/{total_params}**")
        st.download_button(
            "⬇ Download report (.txt)",
            data=report_text.encode(),
            file_name="Wheat_Thesis_Results_Draft.txt",
            mime="text/plain",
        )
    else:
        st.info("Click **🔄 Generate / Refresh report** to build the results draft from current data.")

    # ── Significance summary (always shown) ──────────────────────────────────
    st.divider()
    st.markdown("#### Significance Summary — copy-paste ready sentence")
    _sp = param_options.get(st.session_state.get("sig_sum_sel", ""), None)
    if _sp:
        _summ = generate_sig_summary(_sp)
        if _summ:
            st.text_area("Ready-to-paste result sentence:",
                         _summ, height=120, key="sig_sum_area")
        else:
            st.info("Enter data for this parameter to generate a summary sentence.")
    else:
        st.info("Select a parameter above to generate a significance summary sentence.")
